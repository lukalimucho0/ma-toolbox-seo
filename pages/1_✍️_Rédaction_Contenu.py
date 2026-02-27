import streamlit as st
import requests
import pandas as pd
import time
import random
from bs4 import BeautifulSoup
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import io
import json
from typing import List, Dict, Tuple, Optional
import logging
from dataclasses import dataclass, field

# Charger les variables d'environnement depuis .env
try:
    from openai import OpenAI
except ImportError:
    import openai
import anthropic
import google.generativeai as genai
import trafilatura
from urllib.parse import urlparse
import base64
from datetime import datetime, timedelta
from pathlib import Path
import xml.etree.ElementTree as ET
import glob as glob_module


# Configuration du logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# ═══════════════════════════════════════════════════════════
# DATA CLASSES
# ═══════════════════════════════════════════════════════════

@dataclass
class CompetitorData:
    url: str
    title: str
    meta_description: str
    headings: Dict[str, List[str]]
    position: int
    raw_html: str
    extraction_success: bool
    error_message: str = ""
    relevance_score: float = 0.0
    selected: bool = False
    body_text: str = ""

@dataclass
class HeadingNode:
    level: int
    text: str
    children: List['HeadingNode'] = field(default_factory=list)
    content: str = ""
    sources_used: List[str] = field(default_factory=list)

@dataclass
class SEOBrief:
    target_keyword: str
    optimized_title: str
    optimized_meta_description: str
    headings_structure: str
    country: str
    language: str
    competitors_analyzed: List[str]

@dataclass
class ContentArticle:
    target_keyword: str
    title: str
    meta_description: str
    headings_tree: List[HeadingNode]
    country: str
    language: str
    competitors_analyzed: List[str]
    internal_links_context: str = ""

# ═══════════════════════════════════════════════════════════
# CLIENT BRIEF READER
# ═══════════════════════════════════════════════════════════

class ClientBriefReader:
    """Lit le brief client depuis un fichier uploadé (.docx, .txt, .md)"""
    
    @staticmethod
    def read_uploaded_file(uploaded_file) -> str:
        """Extrait le texte d'un fichier uploadé"""
        if uploaded_file is None:
            return ""
        
        filename = uploaded_file.name.lower()
        
        try:
            if filename.endswith('.docx'):
                return ClientBriefReader._read_docx(uploaded_file)
            elif filename.endswith('.txt') or filename.endswith('.md'):
                return uploaded_file.read().decode('utf-8')
            elif filename.endswith('.doc'):
                return "[Format .doc non supporté. Merci de convertir en .docx ou .txt]"
            else:
                # Tente de lire comme du texte
                return uploaded_file.read().decode('utf-8')
        except Exception as e:
            logger.error(f"Erreur lecture fichier client: {e}")
            return f"[Erreur lecture du fichier: {e}]"
    
    @staticmethod
    def _read_docx(uploaded_file) -> str:
        """Lit un fichier .docx et retourne le texte"""
        doc = Document(io.BytesIO(uploaded_file.read()))
        full_text = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                full_text.append(text)
        return "\n".join(full_text)

# ═══════════════════════════════════════════════════════════
# INTERNAL LINKS PARSER
# ═══════════════════════════════════════════════════════════

class InternalLinksParser:
    """Parse le format URL | ancre de lien pour le maillage interne"""
    
    @staticmethod
    def parse_links_input(raw_input: str) -> List[Dict[str, str]]:
        """Parse le texte brut en liste de {url, anchor}"""
        links = []
        if not raw_input or not raw_input.strip():
            return links
        
        for line in raw_input.strip().split('\n'):
            line = line.strip()
            if not line:
                continue
            
            # Format : URL | ancre de lien
            if '|' in line:
                parts = line.split('|', 1)
                url = parts[0].strip()
                anchor = parts[1].strip()
                if url and anchor:
                    links.append({"url": url, "anchor": anchor})
            # Fallback : URL seule
            elif line.startswith('http'):
                links.append({"url": line, "anchor": ""})
        
        return links
    
    @staticmethod
    def format_for_prompt(links: List[Dict[str, str]]) -> str:
        """Formate les liens pour injection dans le prompt de rédaction"""
        if not links:
            return ""
        
        lines = []
        for link in links:
            if link["anchor"]:
                lines.append(f"- URL: {link['url']} → Ancre à utiliser: \"{link['anchor']}\"")
            else:
                lines.append(f"- URL: {link['url']} (choisis une ancre contextuelle)")
        
        return "\n".join(lines)

# ═══════════════════════════════════════════════════════════
# SITEMAP MANAGER - Gestion des sitemaps clients
# ═══════════════════════════════════════════════════════════

class SitemapManager:
    """Gère le fetch, le parsing et le cache des sitemaps pour chaque client"""

    CLIENTS_DIR = Path(__file__).parent / "clients"

    # Namespaces XML courants dans les sitemaps
    SITEMAP_NS = {
        'sm': 'http://www.sitemaps.org/schemas/sitemap/0.9',
        'image': 'http://www.google.com/schemas/sitemap-image/1.1',
        'xhtml': 'http://www.w3.org/1999/xhtml'
    }

    @classmethod
    def get_available_clients(cls) -> List[Dict]:
        """Retourne la liste des clients configurés"""
        clients = []
        if not cls.CLIENTS_DIR.exists():
            return clients

        for json_file in sorted(cls.CLIENTS_DIR.glob("*.json")):
            try:
                with open(json_file, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    data['_file'] = str(json_file)
                    clients.append(data)
            except Exception as e:
                logger.error(f"Erreur lecture {json_file}: {e}")

        return clients

    @classmethod
    def load_client(cls, client_name: str) -> Optional[Dict]:
        """Charge la configuration d'un client par son nom"""
        for client in cls.get_available_clients():
            if client['name'] == client_name:
                return client
        return None

    @classmethod
    def save_client(cls, client_data: Dict):
        """Sauvegarde la configuration d'un client"""
        file_path = client_data.get('_file')
        if not file_path:
            return

        # Copie sans le champ interne _file
        save_data = {k: v for k, v in client_data.items() if not k.startswith('_')}

        with open(file_path, 'w', encoding='utf-8') as f:
            json.dump(save_data, f, ensure_ascii=False, indent=2)

    @classmethod
    def is_cache_valid(cls, client_data: Dict) -> bool:
        """Vérifie si le cache du sitemap est encore valide"""
        last_fetched = client_data.get('last_fetched')
        if not last_fetched:
            return False

        ttl_hours = client_data.get('cache_ttl_hours', 168)  # 7 jours par défaut
        try:
            fetched_dt = datetime.fromisoformat(last_fetched)
            return datetime.now() - fetched_dt < timedelta(hours=ttl_hours)
        except (ValueError, TypeError):
            return False

    @classmethod
    def fetch_sitemap_xml(cls, url: str, timeout: int = 15) -> Optional[str]:
        """Télécharge le contenu XML d'un sitemap"""
        try:
            headers = {
                'User-Agent': 'Mozilla/5.0 (compatible; SEOContentWriter/1.0)'
            }
            response = requests.get(url, headers=headers, timeout=timeout)
            response.raise_for_status()
            return response.text
        except Exception as e:
            logger.warning(f"Erreur fetch sitemap {url}: {e}")
            return None

    @classmethod
    def parse_sitemap_xml(cls, xml_content: str) -> List[Dict]:
        """Parse le XML d'un sitemap et extrait les URLs"""
        pages = []
        try:
            root = ET.fromstring(xml_content)
            tag = root.tag.lower()

            # Détecter si c'est un sitemap index (contient d'autres sitemaps)
            if 'sitemapindex' in tag:
                # C'est un index : extraire les URLs des sub-sitemaps
                for sitemap in root.findall('.//sm:sitemap/sm:loc', cls.SITEMAP_NS):
                    sub_url = sitemap.text.strip() if sitemap.text else None
                    if sub_url:
                        sub_xml = cls.fetch_sitemap_xml(sub_url)
                        if sub_xml:
                            pages.extend(cls.parse_sitemap_xml(sub_xml))
            else:
                # C'est un sitemap classique : extraire les URLs des pages
                for url_element in root.findall('.//sm:url', cls.SITEMAP_NS):
                    loc = url_element.find('sm:loc', cls.SITEMAP_NS)
                    lastmod = url_element.find('sm:lastmod', cls.SITEMAP_NS)

                    if loc is not None and loc.text:
                        page_url = loc.text.strip()
                        page_data = {
                            'url': page_url,
                            'lastmod': lastmod.text.strip() if lastmod is not None and lastmod.text else None,
                            'title': cls._extract_title_from_url(page_url)
                        }
                        pages.append(page_data)

                # Fallback sans namespace si rien trouvé
                if not pages:
                    for url_element in root.findall('.//{http://www.sitemaps.org/schemas/sitemap/0.9}url'):
                        loc = url_element.find('{http://www.sitemaps.org/schemas/sitemap/0.9}loc')
                        lastmod = url_element.find('{http://www.sitemaps.org/schemas/sitemap/0.9}lastmod')

                        if loc is not None and loc.text:
                            page_url = loc.text.strip()
                            page_data = {
                                'url': page_url,
                                'lastmod': lastmod.text.strip() if lastmod is not None and lastmod.text else None,
                                'title': cls._extract_title_from_url(page_url)
                            }
                            pages.append(page_data)

                # Dernier fallback : sans aucun namespace
                if not pages:
                    for url_element in root.iter():
                        if url_element.tag.endswith('}loc') or url_element.tag == 'loc':
                            if url_element.text and url_element.text.startswith('http'):
                                page_url = url_element.text.strip()
                                pages.append({
                                    'url': page_url,
                                    'lastmod': None,
                                    'title': cls._extract_title_from_url(page_url)
                                })
        except ET.ParseError as e:
            logger.error(f"Erreur parsing XML sitemap: {e}")

        return pages

    @classmethod
    def _extract_title_from_url(cls, url: str) -> str:
        """Extrait un titre lisible depuis l'URL (slug → titre)"""
        parsed = urlparse(url)
        path = parsed.path.strip('/')

        if not path:
            return "Accueil"

        # Prendre le dernier segment du path
        slug = path.split('/')[-1]

        # Nettoyer : retirer extension, remplacer tirets/underscores
        slug = slug.replace('.html', '').replace('.php', '').replace('.htm', '')
        slug = slug.replace('-', ' ').replace('_', ' ')

        # Capitaliser
        return slug.title() if slug else path

    @classmethod
    def fetch_and_cache_sitemap(cls, client_data: Dict, force: bool = False) -> Tuple[List[Dict], str]:
        """
        Fetch les sitemaps d'un client, parse et met en cache.
        Retourne (pages, status_message)
        """
        # Vérifier le cache
        if not force and cls.is_cache_valid(client_data) and client_data.get('pages'):
            age = datetime.now() - datetime.fromisoformat(client_data['last_fetched'])
            hours_ago = int(age.total_seconds() / 3600)
            return client_data['pages'], f"Cache valide (mis à jour il y a {hours_ago}h)"

        sitemap_urls = client_data.get('sitemap_urls', [])
        if not sitemap_urls:
            return [], "Aucun sitemap configuré pour ce client"

        all_pages = []
        errors = []

        for sitemap_url in sitemap_urls:
            xml_content = cls.fetch_sitemap_xml(sitemap_url)
            if xml_content:
                pages = cls.parse_sitemap_xml(xml_content)
                all_pages.extend(pages)
            else:
                errors.append(sitemap_url)

        # Dédupliquer par URL
        seen_urls = set()
        unique_pages = []
        for page in all_pages:
            if page['url'] not in seen_urls:
                seen_urls.add(page['url'])
                unique_pages.append(page)

        # Sauvegarder dans le cache
        client_data['pages'] = unique_pages
        client_data['last_fetched'] = datetime.now().isoformat()
        cls.save_client(client_data)

        status = f"Sitemap mis à jour : {len(unique_pages)} pages trouvées"
        if errors:
            status += f" ({len(errors)} sitemap(s) en erreur)"

        return unique_pages, status

    @classmethod
    def format_sitemap_for_prompt(cls, pages: List[Dict], max_pages: int = 200) -> str:
        """
        Formate les pages du sitemap pour injection dans le prompt de rédaction.
        Le modèle IA choisira les pages pertinentes à mailler.
        """
        if not pages:
            return ""

        # Limiter le nombre de pages si nécessaire
        display_pages = pages[:max_pages]

        lines = []
        for page in display_pages:
            title = page.get('title', '')
            url = page.get('url', '')
            lines.append(f"- {url} | Titre déduit: \"{title}\"")

        if len(pages) > max_pages:
            lines.append(f"\n... et {len(pages) - max_pages} pages supplémentaires non listées.")

        return "\n".join(lines)

    @classmethod
    def get_sitemap_stats(cls, pages: List[Dict]) -> Dict:
        """Retourne des statistiques sur les pages du sitemap"""
        if not pages:
            return {"total": 0}

        # Catégoriser par type d'URL
        stats = {"total": len(pages)}
        for page in pages:
            url = page.get('url', '')
            parsed = urlparse(url)
            path_parts = parsed.path.strip('/').split('/')

            category = path_parts[0] if path_parts and path_parts[0] else "racine"
            stats[category] = stats.get(category, 0) + 1

        return stats


# ═══════════════════════════════════════════════════════════
# DATAFORSEO CONFIG & API
# ═══════════════════════════════════════════════════════════

class DataForSEOConfig:
    SUPPORTED_MARKETS = {
        "🇫🇷 France": {
            "location_code": 2250, "country_code": "FR",
            "languages": {"Français": "fr"}, "google_domain": "google.fr"
        },
        "🇺🇸 États-Unis": {
            "location_code": 2840, "country_code": "US",
            "languages": {"English": "en"}, "google_domain": "google.com"
        },
        "🇬🇧 Royaume-Uni": {
            "location_code": 2826, "country_code": "GB",
            "languages": {"English": "en"}, "google_domain": "google.co.uk"
        },
        "🇩🇪 Allemagne": {
            "location_code": 2276, "country_code": "DE",
            "languages": {"Deutsch": "de", "English": "en"}, "google_domain": "google.de"
        },
        "🇪🇸 Espagne": {
            "location_code": 2724, "country_code": "ES",
            "languages": {"Español": "es", "English": "en"}, "google_domain": "google.es"
        },
        "🇮🇹 Italie": {
            "location_code": 2380, "country_code": "IT",
            "languages": {"Italiano": "it", "English": "en"}, "google_domain": "google.it"
        },
        "🇧🇪 Belgique": {
            "location_code": 2056, "country_code": "BE",
            "languages": {"Français": "fr", "Nederlands": "nl"}, "google_domain": "google.be"
        },
        "🇨🇦 Canada": {
            "location_code": 2124, "country_code": "CA",
            "languages": {"English": "en", "Français": "fr"}, "google_domain": "google.ca"
        }
    }
    
    @classmethod
    def get_market_config(cls, country_name: str, language_name: str) -> Dict:
        if country_name not in cls.SUPPORTED_MARKETS:
            raise ValueError(f"Pays non supporté: {country_name}")
        market = cls.SUPPORTED_MARKETS[country_name]
        if language_name not in market["languages"]:
            raise ValueError(f"Langue non supportée pour {country_name}: {language_name}")
        return {
            "location_code": market["location_code"],
            "language_code": market["languages"][language_name],
            "country_code": market["country_code"],
            "google_domain": market["google_domain"]
        }
    
    @classmethod
    def get_available_languages(cls, country_name: str) -> List[str]:
        if country_name not in cls.SUPPORTED_MARKETS:
            return []
        return list(cls.SUPPORTED_MARKETS[country_name]["languages"].keys())


class DataForSEOAPI:
    def __init__(self, username: str, password: str):
        self.username = username
        self.password = password
        self.base_url = "https://api.dataforseo.com/v3"
        self.session = requests.Session()
        credentials = f"{username}:{password}"
        encoded_credentials = base64.b64encode(credentials.encode()).decode()
        self.session.headers.update({
            "Authorization": f"Basic {encoded_credentials}",
            "Content-Type": "application/json"
        })
    
    def search_serp_live(self, keyword: str, location_code: int, language_code: str, 
                        num_results: int = 10) -> List[Dict]:
        endpoint = f"{self.base_url}/serp/google/organic/live/advanced"
        payload = [{
            "keyword": keyword,
            "location_code": location_code,
            "language_code": language_code,
            "device": "desktop",
            "os": "windows",
            "depth": min(num_results, 100),
            "calculate_rectangles": False
        }]
        try:
            response = self.session.post(endpoint, json=payload)

            # Log HTTP status pour debug
            logger.info(f"DataForSEO HTTP Status: {response.status_code}")

            if response.status_code == 401:
                logger.error("DataForSEO: Authentification échouée (401). Vérifiez email + clé API.")
                raise ValueError("❌ DataForSEO: Authentification échouée (401). Vérifiez que vous utilisez bien votre **email** comme Username et votre **clé API** comme Password (pas le mot de passe du compte). Trouvez-les dans Dashboard → API Settings.")

            if response.status_code == 403:
                logger.error("DataForSEO: Accès refusé (403). Compte désactivé ou crédits épuisés.")
                raise ValueError("❌ DataForSEO: Accès refusé (403). Votre compte est peut-être désactivé ou vos crédits sont épuisés. Vérifiez votre solde sur app.dataforseo.com.")

            response.raise_for_status()
            data = response.json()

            if data.get("status_code") != 20000:
                error_msg = data.get('status_message', 'Erreur inconnue')
                logger.error(f"DataForSEO API Error: {data.get('status_code')} - {error_msg}")

                # Vérifier aussi les erreurs au niveau task
                if data.get("tasks"):
                    task = data["tasks"][0]
                    task_status = task.get("status_code", "?")
                    task_msg = task.get("status_message", "")
                    logger.error(f"DataForSEO Task Status: {task_status} - {task_msg}")
                    raise ValueError(f"❌ DataForSEO erreur API ({task_status}): {task_msg}")

                raise ValueError(f"❌ DataForSEO erreur globale ({data.get('status_code')}): {error_msg}")

            results = []
            if data.get("tasks") and data["tasks"][0].get("result"):
                task = data["tasks"][0]

                # Vérifier le status de la task elle-même
                if task.get("status_code") != 20000:
                    task_msg = task.get("status_message", "Erreur inconnue")
                    logger.error(f"DataForSEO Task Error: {task.get('status_code')} - {task_msg}")
                    raise ValueError(f"❌ DataForSEO task error ({task.get('status_code')}): {task_msg}")

                serp_items = task["result"][0].get("items", [])
                for i, item in enumerate(serp_items, 1):
                    if item.get("type") == "organic" and item.get("url"):
                        results.append({
                            "position": i,
                            "title": item.get("title", ""),
                            "url": item.get("url", ""),
                            "snippet": item.get("description", ""),
                            "meta_description": item.get("description", "")
                        })
                    if len(results) >= num_results:
                        break

            if not results:
                logger.warning(f"DataForSEO: 0 résultats organiques trouvés pour '{keyword}'")
            else:
                logger.info(f"DataForSEO: {len(results)} résultats organiques trouvés")

            return results
        except ValueError:
            raise  # Re-propager les ValueError pour qu'elles s'affichent dans l'UI
        except Exception as e:
            logger.error(f"Erreur DataForSEO API: {e}")
            raise ValueError(f"❌ Erreur de connexion DataForSEO: {e}")

# ═══════════════════════════════════════════════════════════
# EXTRACTION & ANALYSIS
# ═══════════════════════════════════════════════════════════

class TrafilaturaExtractor:
    @staticmethod
    def extract_content_and_headings(url: str) -> CompetitorData:
        try:
            downloaded = trafilatura.fetch_url(url)
            if not downloaded:
                return CompetitorData(
                    url=url, title="", meta_description="", headings={}, 
                    position=0, raw_html="", extraction_success=False,
                    error_message="Impossible de télécharger la page"
                )
            soup = BeautifulSoup(downloaded, 'html.parser')
            title_tag = soup.find('title')
            title = title_tag.get_text().strip() if title_tag else ""
            meta_desc = TrafilaturaExtractor._extract_meta_description(soup)
            headings = TrafilaturaExtractor._extract_headings_from_html(soup)
            body_text = trafilatura.extract(downloaded, include_comments=False, include_tables=True) or ""
            return CompetitorData(
                url=url, title=title, meta_description=meta_desc,
                headings=headings, position=0, raw_html=downloaded[:2000],
                extraction_success=True, error_message="",
                body_text=body_text[:8000]
            )
        except Exception as e:
            logger.error(f"Erreur Trafilatura pour {url}: {e}")
            return CompetitorData(
                url=url, title="", meta_description="", headings={}, 
                position=0, raw_html="", extraction_success=False,
                error_message=str(e)
            )
    
    @staticmethod
    def _extract_meta_description(soup: BeautifulSoup) -> str:
        meta_desc = soup.find('meta', attrs={'name': 'description'})
        if meta_desc and meta_desc.get('content'):
            return meta_desc.get('content', '').strip()
        og_desc = soup.find('meta', attrs={'property': 'og:description'})
        if og_desc and og_desc.get('content'):
            return og_desc.get('content', '').strip()
        return ""
    
    @staticmethod
    def _extract_headings_from_html(soup: BeautifulSoup) -> Dict[str, List[str]]:
        headings = {}
        for level in range(1, 7):
            tag_name = f'h{level}'
            tags = soup.find_all(tag_name)
            if tags:
                heading_texts = []
                for tag in tags:
                    text = tag.get_text().strip()
                    if text and len(text) <= 200:
                        text = re.sub(r'\s+', ' ', text)
                        heading_texts.append(text)
                if heading_texts:
                    headings[f'H{level}'] = heading_texts
        return headings


class RelevanceAnalyzer:
    @staticmethod
    def calculate_keyword_presence_score(title: str, target_keyword: str) -> float:
        if not title or not target_keyword:
            return 0.0
        title_lower = title.lower()
        keyword_lower = target_keyword.lower()
        if keyword_lower in title_lower:
            return 1.0
        keyword_words = keyword_lower.split()
        title_words = title_lower.split()
        matches = sum(1 for word in keyword_words if word in title_words)
        return matches / len(keyword_words) if keyword_words else 0.0
    
    @staticmethod
    def calculate_heading_structure_score(headings: Dict[str, List[str]]) -> float:
        score = 0.0
        if 'H1' in headings and headings['H1']:
            score += 0.4
        if 'H2' in headings and len(headings['H2']) >= 2:
            score += 0.3
        levels_present = [level for level in ['H1', 'H2', 'H3', 'H4'] if level in headings and headings[level]]
        if len(levels_present) >= 3:
            score += 0.3
        return min(score, 1.0)
    
    @classmethod
    def calculate_relevance_score(cls, competitor_data: CompetitorData, target_keyword: str, position: int) -> float:
        position_score = max(0, (11 - position) / 10) if position <= 10 else 0
        keyword_score = cls.calculate_keyword_presence_score(competitor_data.title, target_keyword)
        structure_score = cls.calculate_heading_structure_score(competitor_data.headings)
        return (position_score * 0.4 + keyword_score * 0.4 + structure_score * 0.2)

# ═══════════════════════════════════════════════════════════
# AI ANALYZER
# ═══════════════════════════════════════════════════════════

class AIAnalyzer:
    def __init__(self, ai_provider: str, ai_model: str, api_key: str):
        self.ai_provider = ai_provider.lower()
        self.ai_model = ai_model
        self.api_key = api_key
        if self.ai_provider == 'openai':
            self.openai_client = OpenAI(api_key=api_key)
        elif self.ai_provider == 'claude':
            self.anthropic_client = anthropic.Anthropic(api_key=api_key)
        elif self.ai_provider == 'gemini':
            genai.configure(api_key=api_key)
    
    @staticmethod
    def get_available_models(provider: str) -> List[Dict[str, str]]:
        models = {
            'claude': [
                {'model_id': 'claude-opus-4-6', 'name': 'Claude Opus 4.6',
                 'description': '🏆 Le plus puissant - Dernière génération', 'best_for': 'Analyses complexes, rédaction premium'},
                {'model_id': 'claude-sonnet-4-5-20250929', 'name': 'Claude Sonnet 4.5',
                 'description': '🚀 Excellent rapport qualité/prix - Recommandé', 'best_for': 'Analyses SEO et rédaction'},
                {'model_id': 'claude-haiku-4-5-20251001', 'name': 'Claude Haiku 4.5',
                 'description': '⚡ Ultra rapide et économique', 'best_for': 'Tâches rapides, gros volumes'},
                {'model_id': 'claude-sonnet-4-20250514', 'name': 'Claude Sonnet 4',
                 'description': '⚖️ Performant et équilibré', 'best_for': 'Analyses SEO complexes et rédaction'},
                {'model_id': 'claude-3-5-sonnet-20241022', 'name': 'Claude 3.5 Sonnet',
                 'description': '📌 Version éprouvée et stable', 'best_for': 'Analyses SEO et rédaction'},
                {'model_id': 'claude-3-opus-20240229', 'name': 'Claude 3 Opus',
                 'description': '💪 Puissant pour analyses approfondies', 'best_for': 'Analyses complexes'},
            ],
            'openai': [
                {'model_id': 'gpt-4o', 'name': 'GPT-4o (Omni)',
                 'description': '🚀 Le plus récent et optimisé - Recommandé', 'best_for': 'Analyses SEO avancées'},
                {'model_id': 'gpt-4-turbo', 'name': 'GPT-4 Turbo',
                 'description': '💨 Version optimisée pour la vitesse', 'best_for': 'Analyses rapides'}
            ],
            'gemini': [
                {'model_id': 'gemini-2.5-pro-preview-06-05', 'name': 'Gemini 2.5 Pro',
                 'description': '🚀 Dernière version - Recommandé', 'best_for': 'Analyses SEO complexes'},
                {'model_id': 'gemini-2.5-flash-preview-05-20', 'name': 'Gemini 2.5 Flash',
                 'description': '⚡ Plus rapide et efficace', 'best_for': 'Analyses rapides'}
            ]
        }
        return models.get(provider.lower(), [])
    
    def analyze_with_custom_prompt(self, prompt: str, max_tokens: int = 2000, temperature: float = 0.3) -> str:
        try:
            if self.ai_provider == 'openai':
                response = self.openai_client.chat.completions.create(
                    model=self.ai_model,
                    messages=[
                        {"role": "system", "content": "Tu es un expert SEO senior spécialisé en rédaction de contenu optimisé."},
                        {"role": "user", "content": prompt}
                    ],
                    temperature=temperature, max_tokens=max_tokens
                )
                return response.choices[0].message.content
            elif self.ai_provider == 'claude':
                message = self.anthropic_client.messages.create(
                    model=self.ai_model, max_tokens=max_tokens, temperature=temperature,
                    messages=[{"role": "user", "content": prompt}]
                )
                return message.content[0].text
            elif self.ai_provider == 'gemini':
                model = genai.GenerativeModel(self.ai_model)
                response = model.generate_content(
                    prompt,
                    generation_config=genai.types.GenerationConfig(
                        temperature=temperature, max_output_tokens=max_tokens
                    )
                )
                return response.text
        except Exception as e:
            logger.error(f"Erreur API {self.ai_provider} ({self.ai_model}): {e}")
            return f"Erreur lors de l'analyse: {e}"

# ═══════════════════════════════════════════════════════════
# PROMPT TEMPLATES V2
# ═══════════════════════════════════════════════════════════

class PromptTemplates:
    
    @staticmethod
    def get_structure_prompt(keyword: str, competitors_data: List[CompetitorData]) -> str:
        competitor_structures = ""
        for i, comp in enumerate(competitors_data, 1):
            if comp.extraction_success and comp.headings:
                competitor_structures += f"\n--- CONCURRENT {i} (Position #{comp.position}) ---\n"
                competitor_structures += f"URL: {comp.url}\nTITLE: {comp.title}\n"
                all_headings = []
                for level in ['H1', 'H2', 'H3', 'H4', 'H5', 'H6']:
                    if level in comp.headings:
                        for h in comp.headings[level]:
                            all_headings.append(f"  {level}: {h}")
                competitor_structures += "STRUCTURE:\n" + "\n".join(all_headings) + "\n"
        
        return f"""Tu es un architecte de contenu SEO expert. Ta mission : construire la structure Hn la plus PERTINENTE et RESSERRÉE possible pour un article ciblant la requête « {keyword} ».

## PRINCIPE FONDAMENTAL

Un article resserré et ultra-pertinent surclasse TOUJOURS un article exhaustif mais dilué. Ton objectif n'est PAS de tout couvrir, mais de couvrir l'ESSENTIEL mieux que quiconque.

## MÉTHODOLOGIE STRICTE

Tu dois procéder en 4 phases mentales AVANT de proposer ta structure :

### Phase 1 : Intention de recherche
Formule en UNE phrase l'intention EXACTE de l'utilisateur qui tape « {keyword} ».
Cette phrase est ton FILTRE CENTRAL : chaque heading que tu envisages devra y répondre directement.

### Phase 2 : Identification des thématiques concurrentes
Analyse chaque concurrent ci-dessous et identifie :
- Les thématiques couvertes par la MAJORITÉ des concurrents (patterns communs)
- Les angles éditoriaux dominants (informatif, comparatif, guide pratique, etc.)
- Les thèmes couverts par seulement 1-2 concurrents (souvent périphériques)

### Phase 3 : Filtration par pertinence (ÉTAPE CLÉ)
Pour CHAQUE thématique identifiée, applique ce test :
- **ESSENTIEL** : répond DIRECTEMENT à l'intention de recherche définie en Phase 1 → GARDER
- **PÉRIPHÉRIQUE** : lié au sujet mais pas indispensable pour répondre à la requête → ÉLIMINER

Règle de filtration : si un lecteur qui a tapé « {keyword} » peut se passer de cette info sans que l'article perde sa valeur, c'est PÉRIPHÉRIQUE → élimine-le.

ÉLIMINE systématiquement :
- Les sujets tangentiels que les concurrents traitent par excès de zèle
- Les sous-thématiques qui mériteraient leur propre article séparé
- Les répétitions thématiques sous des angles différents
- Tout ce qui dilue le propos central

### Phase 4 : Construction de la structure focalisée

## CONTRAINTES DE TAILLE (OBLIGATOIRES)

- **4 à 8 H2 maximum** (proportionnel à la complexité du sujet)
- **2 à 4 H3 par H2 maximum**
- **H4** : uniquement si un H3 a réellement besoin d'être subdivisé (rare)
- **15 à 25 headings au total maximum** (H1 compris)
- Chaque heading doit passer le test : « L'utilisateur qui a tapé {keyword} a-t-il BESOIN de cette info pour obtenir sa réponse ? »

## RÈGLES POUR LA STRUCTURE Hn

1. **H1** : Unique. Contient le mot-clé principal. Descriptif et accrocheur. PAS de formule creuse.

2. **H2** : PILIERS thématiques essentiels. Chaque H2 doit :
   - Couvrir un ANGLE DISTINCT et SPÉCIFIQUE du sujet
   - Être suffisamment précis pour que le lecteur sache ce qu'il va trouver
   - NE JAMAIS être un mot générique seul ("Conclusion", "Introduction", "FAQ", "Avis")
   - NE JAMAIS reprendre le mot-clé à l'identique sans valeur ajoutée

3. **H3** : Sous-thématiques concrètes au sein d'un H2. Uniquement si le H2 parent nécessite une subdivision pour rester clair.

4. **INTERDIT** :
   - H2/H3 contenant uniquement "Conclusion", "En résumé", "FAQ", "Questions fréquentes"
   - H2 trop génériques, redondants ou qui s'éloignent de l'intention de recherche
   - Ajouter des sections "pour faire du volume" ou "parce que les concurrents les ont"
   - Dépasser les contraintes de taille ci-dessus

5. **OBLIGATOIRE** :
   - Chaque H2 doit être ANCRÉ sur les données concurrentielles ET passer le filtre de pertinence
   - Préférer la PROFONDEUR sur les sujets essentiels plutôt que la LARGEUR sur tous les sujets

## DONNÉES CONCURRENTIELLES

{competitor_structures}

## FORMAT DE RÉPONSE

Réponds UNIQUEMENT avec la structure Hn (pas d'analyse, pas de commentaires) :

H1: [Ton H1 optimisé]

H2: [Premier pilier thématique essentiel]
  H3: [Sous-thématique concrète]
  H3: [Sous-thématique concrète]

H2: [Deuxième pilier thématique essentiel]
  H3: [Sous-thématique concrète]
  H3: [Sous-thématique concrète]

[etc. — 4 à 8 H2 max, 15-25 headings total max]
"""

    @staticmethod
    def get_title_prompt(keyword: str, competitors_data: List[CompetitorData]) -> str:
        titles_list = ""
        for comp in competitors_data:
            if comp.extraction_success and comp.title:
                titles_list += f"Position #{comp.position}: {comp.title}\n"
        return f"""Tu es un expert SEO. Analyse les titles concurrents ci-dessous pour la requête « {keyword} » et propose UN title optimisé.

TITLES CONCURRENTS :
{titles_list}

CONTRAINTES :
- Maximum 60 caractères (idéal 50-60)
- Contient le mot-clé principal naturellement
- Différenciant par rapport aux concurrents
- Ton attractif qui incite au clic
- Pas de pipe (|) ni de tirets superflus

Réponds UNIQUEMENT avec :
Title: [Ton title ici]
Caractères: [nombre]"""

    @staticmethod
    def get_meta_description_prompt(keyword: str, competitors_data: List[CompetitorData]) -> str:
        metas_list = ""
        for comp in competitors_data:
            if comp.extraction_success and comp.meta_description:
                metas_list += f"Position #{comp.position}: {comp.meta_description}\n"
        return f"""Tu es un expert SEO. Analyse les meta descriptions concurrentes ci-dessous pour la requête « {keyword} » et propose UNE meta description optimisée.

META DESCRIPTIONS CONCURRENTES :
{metas_list}

CONTRAINTES :
- Entre 145-155 caractères
- Contient le mot-clé principal
- Inclut un appel à l'action ou une promesse de valeur
- Différenciante et engageante

Réponds UNIQUEMENT avec :
Meta Description: [Ta meta description ici]
Caractères: [nombre]"""

    @staticmethod
    def get_paragraph_writing_prompt(
        keyword: str,
        current_heading: str,
        heading_level: int,
        parent_heading: str,
        full_structure: str,
        previous_content_summary: str,
        competitor_extracts: str,
        language: str,
        internal_links_formatted: str = "",
        client_brief: str = "",
        already_linked_urls: list = None
    ) -> str:
        """Prompt de rédaction pour UN paragraphe/section spécifique"""
        
        lang_instruction = "en français" if language == "fr" else f"in {language}"
        
        # Bloc brief client
        client_brief_block = ""
        if client_brief.strip():
            client_brief_block = f"""
## BRIEF CLIENT & TON ÉDITORIAL

Le contenu est rédigé pour un client spécifique. Voici les informations à respecter impérativement (identité de marque, ton, vocabulaire, contraintes) :

{client_brief}

Tu DOIS respecter le ton, le vocabulaire et les contraintes décrits dans ce brief. L'article doit sonner comme s'il avait été écrit par l'équipe du client.
"""

        # Bloc maillage interne
        internal_links_block = ""
        if internal_links_formatted.strip():
            # Détecter si on est en mode sitemap (beaucoup de pages) ou liens manuels
            is_sitemap_mode = internal_links_formatted.count('\n') > 10

            if is_sitemap_mode:
                # Bloc URLs déjà utilisées
                already_linked_block = ""
                if already_linked_urls:
                    urls_list = "\n".join(f"- {url}" for url in already_linked_urls)
                    already_linked_block = f"""

⛔ URLS DÉJÀ MAILLÉES DANS L'ARTICLE (NE PAS RÉUTILISER) :
{urls_list}
"""

                internal_links_block = f"""
## MAILLAGE INTERNE INTELLIGENT

Tu as accès au SITEMAP COMPLET du site client. Ta mission est de sélectionner les pages les plus pertinentes à mailler dans cette section.

PROCESSUS DE SÉLECTION :
1. Lis la liste des pages disponibles ci-dessous
2. Identifie les pages dont le sujet est EN RAPPORT DIRECT avec le contenu de cette section
3. Crée une ancre en mot-clé exact et descriptif pour chaque page sélectionnée
4. Intègre le lien au format markdown [ancre](url) de manière naturelle dans le texte

RÈGLES STRICTES :
- Sélectionne uniquement 1 à 2 pages pertinentes pour cette section (pas plus)
- L'ancre DOIT être un mot-clé descriptif et exact (pas "cliquez ici", pas "en savoir plus")
- Le lien doit s'intégrer NATURELLEMENT dans une phrase (pas de rupture syntaxique)
- Ne force JAMAIS un lien si aucune page n'est pertinente pour cette section — il vaut mieux 0 lien qu'un lien forcé
- Ne maille PAS vers la page d'accueil, les pages de contact ou les mentions légales
- ⚠️ RÈGLE ABSOLUE : Chaque URL ne peut être maillée qu'UNE SEULE FOIS dans tout l'article. Si une URL a déjà été utilisée dans une section précédente, tu ne dois PAS la réutiliser. Choisis une AUTRE page du sitemap.
{already_linked_block}
PAGES DISPONIBLES SUR LE SITE :
{internal_links_formatted}
"""
            else:
                # Bloc URLs déjà utilisées (mode manuel)
                already_linked_block_manual = ""
                if already_linked_urls:
                    urls_list_manual = "\n".join(f"- {url}" for url in already_linked_urls)
                    already_linked_block_manual = f"""

⛔ URLS DÉJÀ MAILLÉES (NE PAS RÉUTILISER) :
{urls_list_manual}
"""

                internal_links_block = f"""
## MAILLAGE INTERNE

Tu disposes de liens internes à placer dans le contenu. Pour chaque lien, une ancre de texte est définie.
Intègre le lien au format markdown [ancre](url) à l'endroit le plus naturel et pertinent du texte.

RÈGLES :
- Place le lien UNIQUEMENT si le contexte de ta section est pertinent avec la page cible
- L'ancre doit s'insérer de manière fluide dans la phrase (pas de rupture syntaxique)
- Maximum 1-2 liens par section
- Ne force JAMAIS un lien si ça ne colle pas au contenu de cette section
- ⚠️ Chaque URL ne peut être maillée qu'UNE SEULE FOIS dans tout l'article. Ne réutilise pas une URL déjà placée dans une section précédente.
{already_linked_block_manual}
LIENS DISPONIBLES :
{internal_links_formatted}
"""

        return f"""Tu es un rédacteur web SEO expert. Tu rédiges {lang_instruction} une section d'un article optimisé pour la requête « {keyword} ».
{client_brief_block}
## CONTEXTE DE L'ARTICLE COMPLET

Voici la structure complète de l'article pour comprendre où se situe cette section :
```
{full_structure}
```

## CE QUE TU DOIS RÉDIGER MAINTENANT

Heading : {current_heading} (niveau H{heading_level})
{f"Section parente : {parent_heading}" if parent_heading else ""}

## RÉSUMÉ DU CONTENU DÉJÀ RÉDIGÉ AVANT CETTE SECTION

{previous_content_summary if previous_content_summary else "C'est le début de l'article."}

## EXTRAITS CONCURRENTS PERTINENTS POUR CETTE SECTION

Utilise ces extraits comme SOURCES D'INFORMATION FACTUELLES. Ne copie pas, synthétise et reformule :

{competitor_extracts if competitor_extracts else "Aucun extrait concurrent spécifique disponible. Base-toi sur ton expertise."}
{internal_links_block}
## CONSIGNES DE RÉDACTION

1. **Longueur** : 150-250 mots pour un H2, 100-180 mots pour un H3, 80-120 mots pour un H4
2. **Style** : Fluide, naturel, informatif. Pas de formules creuses ni de remplissage
3. **SEO** : Intègre naturellement des variantes sémantiques de « {keyword} » sans suroptimisation
4. **Factuel** : Appuie-toi sur les sources concurrentes pour apporter des informations concrètes
5. **Transition** : Le texte doit s'enchaîner naturellement avec ce qui précède
6. **PAS DE** : pas de "Dans cette section nous allons voir...", pas de conclusion de section, pas de répétition du heading en première phrase
7. **Format** : Rédige uniquement le contenu textuel (pas le heading). Paragraphes + listes à puces si pertinent

## RÉPONSE

Rédige directement le contenu de cette section :"""

# ═══════════════════════════════════════════════════════════
# HEADING STRUCTURE PARSER
# ═══════════════════════════════════════════════════════════

class HeadingParser:
    @staticmethod
    def parse_structure_text(text: str) -> List[HeadingNode]:
        lines = text.strip().split('\n')
        nodes = []
        for line in lines:
            line = line.strip()
            if not line:
                continue
            match = re.match(r'^(?:\*\s*)?(?:\*\*)?H(\d)(?:\*\*)?[\s:]*[:.]?\s*(.+?)(?:\*\*)?$', line, re.IGNORECASE)
            if not match:
                match = re.match(r'^\s*H(\d)\s*[:.-]\s*(.+)$', line, re.IGNORECASE)
            if not match:
                match = re.match(r'^\s*[-*]\s*(?:\*\*)?H(\d)(?:\*\*)?[\s:]*[:.]?\s*(.+?)(?:\*\*)?$', line, re.IGNORECASE)
            if match:
                level = int(match.group(1))
                text_content = match.group(2).strip().strip('*').strip()
                nodes.append(HeadingNode(level=level, text=text_content))
        return nodes
    
    @staticmethod
    def nodes_to_text(nodes: List[HeadingNode]) -> str:
        lines = []
        for node in nodes:
            indent = "  " * (node.level - 1)
            lines.append(f"{indent}H{node.level}: {node.text}")
        return "\n".join(lines)
    
    @staticmethod
    def find_relevant_competitor_content(heading_text: str, competitors_data: List[CompetitorData], keyword: str) -> str:
        relevant_extracts = []
        heading_words = set(heading_text.lower().split())
        stop_words = {'le', 'la', 'les', 'de', 'du', 'des', 'un', 'une', 'et', 'ou', 'en', 'à', 'pour',
                      'the', 'a', 'an', 'of', 'in', 'to', 'and', 'or', 'for', 'with', 'on', 'at',
                      'que', 'qui', 'est', 'sont', 'dans', 'par', 'sur', 'ce', 'cette', 'ces',
                      'comment', 'quoi', 'quel', 'quelle', 'quels', 'quelles', 'votre', 'vos',
                      'notre', 'nos', 'son', 'sa', 'ses', 'leur', 'leurs', 'tout', 'tous', 'toute', 'toutes'}
        heading_keywords = heading_words - stop_words
        
        for comp in competitors_data:
            if not comp.extraction_success or not comp.body_text:
                continue
            paragraphs = comp.body_text.split('\n')
            for para in paragraphs:
                para = para.strip()
                if len(para) < 50:
                    continue
                para_lower = para.lower()
                match_count = sum(1 for w in heading_keywords if w in para_lower)
                if match_count >= max(1, len(heading_keywords) // 2):
                    relevant_extracts.append({'text': para[:500], 'source': comp.url, 'score': match_count})
        
        relevant_extracts.sort(key=lambda x: x['score'], reverse=True)
        top_extracts = relevant_extracts[:5]
        if not top_extracts:
            return ""
        result = ""
        for ext in top_extracts:
            domain = urlparse(ext['source']).netloc
            result += f"[Source: {domain}] {ext['text']}\n\n"
        return result.strip()

# ═══════════════════════════════════════════════════════════
# WORD GENERATOR V2
# ═══════════════════════════════════════════════════════════

class WordGenerator:
    
    @staticmethod
    def create_structure_only_document(keyword: str, structure_text: str, title: str = "", meta_desc: str = "", country: str = "", language: str = "") -> io.BytesIO:
        """Génère un .docx contenant uniquement la structure Hn validée"""
        doc = Document()
        doc_title = doc.add_heading(f'Structure Hn - {keyword}', 0)
        doc_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_paragraph(f'Date: {time.strftime("%d/%m/%Y")} | Marché: {country} | Langue: {language}').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_paragraph('')
        
        if title:
            doc.add_heading('Balise Title', level=2)
            p = doc.add_paragraph(title)
            p.style = 'Intense Quote'
        if meta_desc:
            doc.add_heading('Meta Description', level=2)
            p = doc.add_paragraph(meta_desc)
            p.style = 'Intense Quote'
        
        doc.add_paragraph('')
        doc.add_heading('Structure Hn', level=1)
        
        # Parse et affiche avec indentation
        nodes = HeadingParser.parse_structure_text(structure_text)
        for node in nodes:
            indent = "    " * (node.level - 1)
            level_text = f"H{node.level}: {node.text}"
            p = doc.add_paragraph(f"{indent}{level_text}")
            if node.level == 1:
                run = p.runs[0]
                run.bold = True
                run.font.size = Pt(14)
            elif node.level == 2:
                run = p.runs[0]
                run.bold = True
                run.font.size = Pt(12)
        
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
    
    @staticmethod
    def create_seo_brief_document(seo_brief: SEOBrief) -> io.BytesIO:
        doc = Document()
        title = doc.add_heading(f'Brief SEO - {seo_brief.target_keyword}', 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_paragraph(f'Marché: {seo_brief.country} | Langue: {seo_brief.language}').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_paragraph(f'Date: {time.strftime("%d/%m/%Y")}').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_page_break()
        doc.add_heading('1. Optimisation On-Page', level=1)
        doc.add_heading('Balise Title', level=2)
        p = doc.add_paragraph(seo_brief.optimized_title)
        p.style = 'Intense Quote'
        doc.add_paragraph(f'Caractères: {len(seo_brief.optimized_title)}')
        doc.add_heading('Meta Description', level=2)
        p = doc.add_paragraph(seo_brief.optimized_meta_description)
        p.style = 'Intense Quote'
        doc.add_paragraph(f'Caractères: {len(seo_brief.optimized_meta_description)}')
        doc.add_heading('2. Structure des Headings', level=1)
        doc.add_paragraph(seo_brief.headings_structure)
        if seo_brief.competitors_analyzed:
            doc.add_heading('3. Sources Analysées', level=1)
            for i, source in enumerate(seo_brief.competitors_analyzed, 1):
                doc.add_paragraph(f'{i}. {source}')
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
    
    @staticmethod
    def _clean_seo_field(raw_text: str, field_type: str = 'title') -> str:
        """Nettoie la réponse IA pour extraire uniquement la valeur du Title ou Meta Description"""
        if not raw_text:
            return ""
        # Prendre seulement la première ligne pertinente
        lines = raw_text.strip().split('\n')
        for line in lines:
            line = line.strip()
            # Retirer les préfixes courants de l'IA
            for prefix in ['Title:', 'Title :', 'Meta Description:', 'Meta Description :',
                           'Meta description:', 'Meta description :', 'Caractères:', 'Caractères :']:
                if line.lower().startswith(prefix.lower()):
                    line = line[len(prefix):].strip()
                    break
            # Ignorer les lignes "Caractères: XX"
            if re.match(r'^Caract[eè]res\s*:\s*\d+', line, re.IGNORECASE):
                continue
            if line and not line.startswith('Caractère'):
                return line
        return raw_text.strip().split('\n')[0]

    @staticmethod
    def _set_heading_black(heading_paragraph):
        """Force tous les runs d'un heading en noir"""
        for run in heading_paragraph.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)

    @staticmethod
    def _add_formatted_text(paragraph, text: str):
        """Convertit le markdown inline en mise en forme Word native.
        Gère : **gras**, *italique*, [ancre](url)"""
        # Pattern pour capturer **gras**, *italique*, et [liens](url)
        pattern = r'(\*\*(.+?)\*\*|\*(.+?)\*|\[(.+?)\]\((.+?)\))'
        last_end = 0
        for match in re.finditer(pattern, text):
            # Texte avant le match
            before = text[last_end:match.start()]
            if before:
                paragraph.add_run(before)

            if match.group(2):  # **gras**
                run = paragraph.add_run(match.group(2))
                run.bold = True
            elif match.group(3):  # *italique*
                run = paragraph.add_run(match.group(3))
                run.italic = True
            elif match.group(4) and match.group(5):  # [ancre](url)
                run = paragraph.add_run(match.group(4))
                run.underline = True
                run.font.color.rgb = RGBColor(0, 102, 204)

            last_end = match.end()

        # Texte restant après le dernier match
        remaining = text[last_end:]
        if remaining:
            paragraph.add_run(remaining)

    @staticmethod
    def _add_content_to_doc(doc: Document, content: str):
        """Ajoute du contenu en convertissant le markdown en mise en forme Word"""
        paragraphs = content.strip().split('\n\n')
        for para_text in paragraphs:
            para_text = para_text.strip()
            if not para_text:
                continue

            # Listes à puces (lignes commençant par - ou *)
            lines = para_text.split('\n')
            is_list = all(
                l.strip().startswith('- ') or l.strip().startswith('* ') or not l.strip()
                for l in lines if l.strip()
            )

            if is_list:
                for line in lines:
                    line = line.strip().lstrip('-*').strip()
                    if line:
                        p = doc.add_paragraph(style='List Bullet')
                        WordGenerator._add_formatted_text(p, line)
            else:
                # Paragraphe normal - gérer les sauts de ligne simples
                full_text = ' '.join(l.strip() for l in lines if l.strip())
                p = doc.add_paragraph()
                WordGenerator._add_formatted_text(p, full_text)

    @staticmethod
    def create_full_article_document(article: ContentArticle) -> io.BytesIO:
        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)

        # Forcer tous les styles de heading en noir
        for i in range(1, 5):
            heading_style = doc.styles[f'Heading {i}']
            heading_style.font.color.rgb = RGBColor(0, 0, 0)

        # Nettoyer le title et la meta description (retirer les prefixes IA)
        clean_title = WordGenerator._clean_seo_field(article.title, 'title')
        clean_meta = WordGenerator._clean_seo_field(article.meta_description, 'meta')

        # Commencer directement par Title et Meta Description
        p_title = doc.add_paragraph()
        run_label = p_title.add_run('Title : ')
        run_label.bold = True
        run_label.font.size = Pt(12)
        run_value = p_title.add_run(clean_title)
        run_value.font.size = Pt(12)

        p_meta = doc.add_paragraph()
        run_label = p_meta.add_run('Meta description : ')
        run_label.bold = True
        run_label.font.size = Pt(12)
        run_value = p_meta.add_run(clean_meta)
        run_value.font.size = Pt(12)

        doc.add_paragraph('')  # Espacement

        for node in article.headings_tree:
            WordGenerator._add_heading_node_to_doc(doc, node)

        if article.competitors_analyzed:
            doc.add_page_break()
            h = doc.add_heading('Sources analysées', level=1)
            WordGenerator._set_heading_black(h)
            for i, source in enumerate(article.competitors_analyzed, 1):
                doc.add_paragraph(f'{i}. {source}')

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    @staticmethod
    def _add_heading_node_to_doc(doc: Document, node: HeadingNode):
        h = doc.add_heading(node.text, level=min(node.level, 4))
        WordGenerator._set_heading_black(h)
        if node.content:
            WordGenerator._add_content_to_doc(doc, node.content)
        for child in node.children:
            WordGenerator._add_heading_node_to_doc(doc, child)

# ═══════════════════════════════════════════════════════════
# SEO BRIEF GENERATOR
# ═══════════════════════════════════════════════════════════

class SEOBriefGenerator:
    def __init__(self):
        self.dataforseo_api = None
        self.extractor = TrafilaturaExtractor()
        self.analyzer = RelevanceAnalyzer()
        self.ai_analyzer = None
        
    def setup_apis(self, dataforseo_username, dataforseo_password, ai_provider, ai_model, ai_api_key):
        self.dataforseo_api = DataForSEOAPI(dataforseo_username, dataforseo_password)
        self.ai_analyzer = AIAnalyzer(ai_provider, ai_model, ai_api_key)
    
    def search_and_extract_competitors(self, target_keyword, country, language, num_results=10):
        if not self.dataforseo_api:
            raise ValueError("DataForSEO API non configurée")
        market_config = DataForSEOConfig.get_market_config(country, language)
        search_results = self.dataforseo_api.search_serp_live(
            target_keyword, market_config["location_code"], market_config["language_code"], num_results
        )
        competitors = []
        for result in search_results:
            competitor_data = self.extractor.extract_content_and_headings(result['url'])
            competitor_data.position = result['position']
            if not competitor_data.extraction_success:
                competitor_data.title = result.get('title', '')
                competitor_data.meta_description = result.get('meta_description', '')
            competitor_data.relevance_score = self.analyzer.calculate_relevance_score(
                competitor_data, target_keyword, result['position']
            )
            competitors.append(competitor_data)
        return competitors
    
    def auto_select_competitors(self, competitors, max_competitors=5):
        sorted_competitors = sorted(competitors, key=lambda x: x.relevance_score, reverse=True)
        selected = sorted_competitors[:max_competitors]
        for c in selected:
            c.selected = True
        return selected

# ═══════════════════════════════════════════════════════════
# WRITING ENGINE (extracted for auto/manual modes)
# ═══════════════════════════════════════════════════════════

def _extract_linked_urls(content: str) -> set:
    """Extrait toutes les URLs de liens markdown [texte](url) d'un contenu"""
    if not content:
        return set()
    return set(re.findall(r'\[[^\]]+\]\((https?://[^)]+)\)', content))


def run_writing_engine(
    nodes: List[HeadingNode],
    target_keyword: str,
    selected_competitors: List[CompetitorData],
    ai_analyzer: AIAnalyzer,
    language_code: str,
    internal_links_formatted: str,
    client_brief: str,
    writing_prompt_template: str = "",
    progress_callback=None,
    status_callback=None
) -> List[HeadingNode]:
    """Moteur de rédaction : rédige chaque heading un par un"""

    full_structure_text = HeadingParser.nodes_to_text(nodes)
    content_written = []
    previous_summary = ""
    already_linked_urls = set()  # Tracking des URLs déjà maillées

    for idx, node in enumerate(nodes):
        if progress_callback:
            progress_callback((idx + 1) / len(nodes))
        if status_callback:
            status_callback(f"**✍️ Rédaction [{idx+1}/{len(nodes)}]** : H{node.level} - {node.text}")

        # Skip H1
        if node.level == 1:
            node.content = ""
            content_written.append(node)
            continue

        # Parent heading
        parent_heading = ""
        for prev_node in reversed(content_written):
            if prev_node.level < node.level:
                parent_heading = f"H{prev_node.level}: {prev_node.text}"
                break

        # Extraits concurrents
        competitor_extracts = HeadingParser.find_relevant_competitor_content(
            node.text, selected_competitors, target_keyword
        )

        # Résumé du contenu précédent
        if content_written:
            recent_content = []
            char_count = 0
            for prev_node in reversed(content_written):
                if prev_node.content and char_count < 500:
                    recent_content.insert(0, f"[H{prev_node.level}: {prev_node.text}] {prev_node.content[:200]}...")
                    char_count += len(prev_node.content[:200])
            previous_summary = "\n".join(recent_content[-3:])

        # Génération du prompt (avec URLs déjà utilisées)
        writing_prompt = PromptTemplates.get_paragraph_writing_prompt(
            keyword=target_keyword,
            current_heading=node.text,
            heading_level=node.level,
            parent_heading=parent_heading,
            full_structure=full_structure_text,
            previous_content_summary=previous_summary,
            competitor_extracts=competitor_extracts,
            language=language_code,
            internal_links_formatted=internal_links_formatted,
            client_brief=client_brief,
            already_linked_urls=list(already_linked_urls) if already_linked_urls else None
        )
        
        # Si un template custom a été fourni, l'utiliser
        if writing_prompt_template.strip():
            writing_prompt = writing_prompt_template.replace("{HEADING}", node.text)\
                .replace("{HEADING_LEVEL}", str(node.level))\
                .replace("{KEYWORD}", target_keyword)\
                .replace("{PARENT_HEADING}", parent_heading)\
                .replace("{FULL_STRUCTURE}", full_structure_text)\
                .replace("{PREVIOUS_SUMMARY}", previous_summary)\
                .replace("{COMPETITOR_EXTRACTS}", competitor_extracts)\
                .replace("{INTERNAL_LINKS}", internal_links_formatted)\
                .replace("{CLIENT_BRIEF}", client_brief)\
                .replace("{LANGUAGE}", language_code)
        
        written_content = ai_analyzer.analyze_with_custom_prompt(
            writing_prompt, max_tokens=1500, temperature=0.5
        )

        node.content = written_content
        content_written.append(node)

        # Tracker les URLs maillées dans cette section pour ne pas les réutiliser
        new_urls = _extract_linked_urls(written_content)
        already_linked_urls.update(new_urls)

        time.sleep(0.5)

    return content_written

# ═══════════════════════════════════════════════════════════
# STREAMLIT UI
# ═══════════════════════════════════════════════════════════

def main():
    st.set_page_config(page_title="Rédaction Contenu | Ma Toolbox SEO", page_icon="✍️", layout="wide")
    
    st.title("✍️ SEO Content Writer Pro V2")
    st.markdown("*Analyse concurrentielle → Structure Hn → Rédaction paragraphe par paragraphe*")
    
    # ─── Sidebar ───
    with st.sidebar:
        st.header("⚙️ Configuration")
        
        # Mode
        st.subheader("⚡ Mode")
        execution_mode = st.radio(
            "Mode d'exécution",
            ["🖐️ Manuel (étape par étape)", "🤖 Automatique (tout en un)"],
            index=0,
            help="Manuel : vous validez chaque étape. Automatique : tout s'enchaîne automatiquement."
        )
        auto_mode = "Automatique" in execution_mode
        
        st.divider()
        
        # APIs
        st.subheader("🔍 DataForSEO API")
        dataforseo_username = st.text_input("Username DataForSEO", value=st.secrets.get("DATAFORSEO_USERNAME", ""), type="password")
        dataforseo_password = st.text_input("Password DataForSEO", value=st.secrets.get("DATAFORSEO_PASSWORD", ""), type="password")
        
        st.subheader("🤖 Intelligence Artificielle")
        ai_provider = st.selectbox("Fournisseur IA", ["Claude", "OpenAI", "Gemini"])
        available_models = AIAnalyzer.get_available_models(ai_provider.lower())
        if available_models:
            model_options = [m['name'] for m in available_models]
            selected_model_index = st.selectbox(
                f"Modèle {ai_provider}", range(len(model_options)),
                format_func=lambda x: model_options[x]
            )
            selected_model_info = available_models[selected_model_index]
            st.info(f"**{selected_model_info['description']}**\n\n*Idéal pour :* {selected_model_info['best_for']}")
            ai_model = selected_model_info['model_id']
        # Clé API selon le provider sélectionné
        ai_api_key_env = {
            "Claude": st.secrets.get("ANTHROPIC_API_KEY", ""),
            "OpenAI": st.secrets.get("OPENAI_API_KEY", ""),
            "Gemini": st.secrets.get("GEMINI_API_KEY", "")
        }.get(ai_provider, "")
        ai_api_key = st.text_input(f"Clé API {ai_provider}", value=ai_api_key_env, type="password")
        
        st.divider()
        
        # Brief client
        st.subheader("📄 Brief client (optionnel)")
        client_file = st.file_uploader(
            "Charger un brief client (.docx, .txt)",
            type=["docx", "txt", "md"],
            help="Document contenant : identité du client, ton éditorial, vocabulaire à utiliser, contraintes spécifiques..."
        )
        client_brief_text = ""
        if client_file:
            client_brief_text = ClientBriefReader.read_uploaded_file(client_file)
            st.success(f"✅ Brief chargé ({len(client_brief_text)} caractères)")
            with st.expander("📄 Aperçu du brief"):
                st.text(client_brief_text[:1000] + ("..." if len(client_brief_text) > 1000 else ""))
        
        # Complément manuel au brief
        client_brief_manual = st.text_area(
            "Complément brief / ton éditorial",
            placeholder="Ex: Tutoiement, ton décontracté mais expert. Le client est une marque de literie premium...",
            height=100,
            help="Ajoutez des instructions de ton et de style qui seront injectées dans chaque prompt de rédaction."
        )
        
        # Combinaison brief fichier + manuel
        full_client_brief = ""
        if client_brief_text:
            full_client_brief += client_brief_text + "\n\n"
        if client_brief_manual:
            full_client_brief += client_brief_manual
        
        st.divider()

        # ─── Sélection client & Maillage interne ───
        st.subheader("🏢 Client & Maillage interne")

        # Charger les clients disponibles
        available_clients = SitemapManager.get_available_clients()
        client_names = ["-- Aucun client --"] + [c['name'] for c in available_clients]

        selected_client_name = st.selectbox(
            "Client",
            client_names,
            index=0,
            help="Sélectionnez un client pour charger automatiquement son sitemap et activer le maillage interne intelligent."
        )

        # Variables pour le maillage
        sitemap_pages = []
        sitemap_formatted = ""
        internal_links_input = ""  # Compatibilité avec l'export - sera enrichi ci-dessous
        parsed_links = []

        if selected_client_name != "-- Aucun client --":
            client_data = SitemapManager.load_client(selected_client_name)

            if client_data:
                has_sitemaps = bool(client_data.get('sitemap_urls'))

                if has_sitemaps:
                    # Boutons d'action sitemap
                    col_fetch, col_refresh = st.columns([2, 1])

                    with col_fetch:
                        # Auto-fetch si cache valide
                        if SitemapManager.is_cache_valid(client_data) and client_data.get('pages'):
                            sitemap_pages = client_data['pages']
                            age = datetime.now() - datetime.fromisoformat(client_data['last_fetched'])
                            hours_ago = int(age.total_seconds() / 3600)
                            st.success(f"✅ {len(sitemap_pages)} pages (cache {hours_ago}h)")
                        else:
                            st.info("📡 Sitemap non chargé")

                    with col_refresh:
                        if st.button("🔄", help="Rafraîchir le sitemap", key="refresh_sitemap"):
                            with st.spinner("Chargement..."):
                                sitemap_pages, status_msg = SitemapManager.fetch_and_cache_sitemap(client_data, force=True)
                                st.success(f"✅ {len(sitemap_pages)} pages")

                    # Charger le sitemap automatiquement si pas en cache
                    if not sitemap_pages and has_sitemaps:
                        if st.button("📡 Charger le sitemap", type="primary", key="load_sitemap"):
                            with st.spinner(f"Chargement du sitemap {selected_client_name}..."):
                                sitemap_pages, status_msg = SitemapManager.fetch_and_cache_sitemap(client_data, force=False)
                                st.info(status_msg)

                    # Affichage des pages du sitemap
                    if sitemap_pages:
                        sitemap_formatted = SitemapManager.format_sitemap_for_prompt(sitemap_pages)

                        stats = SitemapManager.get_sitemap_stats(sitemap_pages)

                        with st.expander(f"📄 Pages du sitemap ({stats['total']})", expanded=False):
                            # Affichage des catégories
                            for key, val in stats.items():
                                if key != 'total':
                                    st.caption(f"  📁 /{key}/ → {val} pages")

                            st.divider()

                            # Liste des pages (max 50 affichées)
                            for i, page in enumerate(sitemap_pages[:50]):
                                parsed_url = urlparse(page['url'])
                                st.caption(f"  {parsed_url.path}")
                            if len(sitemap_pages) > 50:
                                st.caption(f"  ... et {len(sitemap_pages) - 50} pages de plus")

                        st.caption("💡 L'IA choisira automatiquement les pages pertinentes à mailler dans l'article.")
                else:
                    st.warning(f"⚠️ Pas de sitemap configuré pour {selected_client_name}")

                # Option de liens manuels complémentaires
                with st.expander("➕ Liens manuels complémentaires", expanded=False):
                    manual_links_input = st.text_area(
                        "Liens manuels (URL | ancre)",
                        placeholder="https://monsite.com/page | ancre exacte",
                        height=80,
                        help="Ajoutez des liens spécifiques en complément du sitemap.",
                        key="manual_links_complement"
                    )
                    manual_parsed = InternalLinksParser.parse_links_input(manual_links_input)
                    if manual_parsed:
                        parsed_links = manual_parsed
                        st.caption(f"✅ {len(manual_parsed)} lien(s) manuel(s)")
        else:
            # Mode sans client : ancien comportement (liens manuels uniquement)
            internal_links_input_raw = st.text_area(
                "Liens internes manuels (URL | ancre)",
                placeholder="https://monsite.com/matelas-latex | matelas en latex naturel\nhttps://monsite.com/guide-sommeil | guide complet du sommeil",
                height=120,
                help="Format : URL | ancre de texte (une par ligne). Sans client sélectionné, utilisez ce champ.",
                key="manual_links_no_client"
            )
            parsed_links = InternalLinksParser.parse_links_input(internal_links_input_raw)
            if parsed_links:
                st.caption(f"✅ {len(parsed_links)} lien(s) détecté(s)")
                for link in parsed_links:
                    anchor_display = f'"{link["anchor"]}"' if link["anchor"] else "(ancre auto)"
                    st.caption(f"  → {anchor_display} → {urlparse(link['url']).path}")
    
    # ─── Main Interface ───
    col1, col2 = st.columns([2, 1])
    
    with col1:
        st.header("🎯 Configuration de recherche")
        countries = list(DataForSEOConfig.SUPPORTED_MARKETS.keys())
        selected_country = st.selectbox("Pays", countries)
        available_languages = DataForSEOConfig.get_available_languages(selected_country)
        selected_language = st.selectbox("Langue", available_languages)
        target_keyword = st.text_input("Mot-clé principal", placeholder="Exemple: comment choisir un lit coffre")
        num_results = st.slider("Nombre de résultats SERP à analyser", 5, 15, 10)
    
    with col2:
        if selected_country and selected_language:
            config = DataForSEOConfig.get_market_config(selected_country, selected_language)
            # Résumé maillage
            if sitemap_pages:
                maillage_info = f"🏢 {selected_client_name} ({len(sitemap_pages)} pages)"
            elif parsed_links:
                maillage_info = f"🔗 {len(parsed_links)} lien(s) manuel(s)"
            else:
                maillage_info = "🔗 Pas de maillage"

            st.info(f"""**Configuration :**
📍 {selected_country}
🗣️ {selected_language}
🔗 {config['google_domain']}
{"📄 Brief client chargé" if full_client_brief else "📄 Pas de brief client"}
{maillage_info}
{"🤖 Mode automatique" if auto_mode else "🖐️ Mode manuel"}""")
    
    # ─── Init ───
    if 'generator' not in st.session_state:
        st.session_state.generator = SEOBriefGenerator()
    
    apis_configured = all([dataforseo_username, dataforseo_password, ai_api_key])
    if not apis_configured:
        st.warning("⚠️ Veuillez configurer toutes les APIs dans la barre latérale")
        return
    
    try:
        st.session_state.generator.setup_apis(
            dataforseo_username, dataforseo_password, ai_provider.lower(), ai_model, ai_api_key
        )
    except Exception as e:
        st.error(f"Erreur configuration API: {e}")
        return
    
    # Session state defaults
    state_defaults = {
        'competitors_data': None, 'search_completed': False, 'selected_competitors': None,
        'structure_result': None, 'title_result': None, 'meta_result': None,
        'validated_structure': None, 'validated_nodes': None,
        'article_content': None, 'writing_done': False
    }
    for key, default in state_defaults.items():
        if key not in st.session_state:
            st.session_state[key] = default
    
    language_code = DataForSEOConfig.get_market_config(selected_country, selected_language)["language_code"]

    # Construire le contexte de maillage interne
    # Priorité : sitemap client > liens manuels
    internal_links_formatted = ""

    if sitemap_formatted:
        # Mode sitemap intelligent : l'IA reçoit toutes les pages et choisit
        internal_links_formatted = sitemap_formatted

    # Ajouter les liens manuels complémentaires
    manual_links_formatted = InternalLinksParser.format_for_prompt(parsed_links)
    if manual_links_formatted:
        if internal_links_formatted:
            internal_links_formatted += "\n\nLIENS MANUELS PRIORITAIRES (à placer obligatoirement) :\n" + manual_links_formatted
        else:
            internal_links_formatted = manual_links_formatted

    # Mettre à jour internal_links_input pour la compatibilité export
    if sitemap_pages:
        internal_links_input = f"[Sitemap {selected_client_name} : {len(sitemap_pages)} pages]"
    elif parsed_links:
        internal_links_input = "\n".join(f"{l['url']} | {l['anchor']}" for l in parsed_links)

    # ═══════════════════════════════════════════════════════
    # MODE AUTOMATIQUE
    # ═══════════════════════════════════════════════════════
    
    if auto_mode:
        st.header("🤖 Mode Automatique")
        st.markdown("Toutes les étapes vont s'enchaîner automatiquement : recherche → extraction → structure Hn → Title/Meta → rédaction complète.")
        
        if st.button("🚀 Lancer la génération complète", type="primary", disabled=not target_keyword):
            
            # Étape 1 : Recherche
            with st.spinner("🔍 1/6 - Recherche SERP et extraction..."):
                try:
                    competitors_data = st.session_state.generator.search_and_extract_competitors(
                        target_keyword, selected_country, selected_language, num_results
                    )
                except Exception as e:
                    st.error(str(e))
                    return
                if not competitors_data:
                    st.error("❌ Aucun concurrent trouvé. Vérifiez vos identifiants DataForSEO (email + clé API, pas le mot de passe du compte).")
                    return
                st.session_state.competitors_data = competitors_data
                st.session_state.search_completed = True
                st.success(f"✅ 1/6 - {len(competitors_data)} concurrents trouvés")
            
            # Étape 2 : Sélection auto
            with st.spinner("🤖 2/6 - Sélection automatique..."):
                selected = st.session_state.generator.auto_select_competitors(competitors_data, max_competitors=5)
                st.session_state.selected_competitors = selected
                st.success(f"✅ 2/6 - {len(selected)} concurrents sélectionnés")
            
            # Étape 3 : Structure Hn
            with st.spinner("🏗️ 3/6 - Génération de la structure Hn..."):
                structure_prompt = PromptTemplates.get_structure_prompt(target_keyword, selected)
                structure_result = st.session_state.generator.ai_analyzer.analyze_with_custom_prompt(structure_prompt, max_tokens=3000)
                st.session_state.structure_result = structure_result
                st.session_state.validated_structure = structure_result
                parsed_nodes = HeadingParser.parse_structure_text(structure_result)
                st.session_state.validated_nodes = parsed_nodes
                st.success(f"✅ 3/6 - Structure Hn générée ({len(parsed_nodes)} headings)")
            
            # Étape 4 : Title
            with st.spinner("🏷️ 4/6 - Génération du Title..."):
                title_result = st.session_state.generator.ai_analyzer.analyze_with_custom_prompt(
                    PromptTemplates.get_title_prompt(target_keyword, selected), max_tokens=200
                )
                st.session_state.title_result = title_result
                st.success("✅ 4/6 - Title généré")
            
            # Étape 5 : Meta Description
            with st.spinner("📝 5/6 - Génération de la Meta Description..."):
                meta_result = st.session_state.generator.ai_analyzer.analyze_with_custom_prompt(
                    PromptTemplates.get_meta_description_prompt(target_keyword, selected), max_tokens=300
                )
                st.session_state.meta_result = meta_result
                st.success("✅ 5/6 - Meta Description générée")
            
            # Étape 6 : Rédaction
            st.markdown("### ✍️ 6/6 - Rédaction de l'article")
            progress_bar = st.progress(0)
            status_text = st.empty()
            
            article_content = run_writing_engine(
                nodes=parsed_nodes,
                target_keyword=target_keyword,
                selected_competitors=selected,
                ai_analyzer=st.session_state.generator.ai_analyzer,
                language_code=language_code,
                internal_links_formatted=internal_links_formatted,
                client_brief=full_client_brief,
                progress_callback=lambda p: progress_bar.progress(p),
                status_callback=lambda s: status_text.markdown(s)
            )
            
            progress_bar.progress(1.0)
            status_text.markdown("**✅ Rédaction terminée !**")
            st.session_state.article_content = article_content
            st.session_state.writing_done = True
            st.success("🎉 Génération automatique terminée !")
        
        # Affichage résultats auto
        if st.session_state.writing_done and st.session_state.article_content:
            _display_article_and_exports(
                st.session_state.article_content, target_keyword,
                st.session_state.title_result, st.session_state.meta_result,
                st.session_state.validated_structure, selected_country, selected_language,
                st.session_state.selected_competitors, internal_links_input
            )
    
    # ═══════════════════════════════════════════════════════
    # MODE MANUEL
    # ═══════════════════════════════════════════════════════
    
    else:
        # ÉTAPE 1 : RECHERCHE
        st.header("🔍 Étape 1 : Recherche et extraction des concurrents")
        
        if st.button("🔍 Rechercher les concurrents", type="primary", disabled=not target_keyword):
            with st.spinner("🔍 Recherche SERP et extraction du contenu..."):
                try:
                    competitors_data = st.session_state.generator.search_and_extract_competitors(
                        target_keyword, selected_country, selected_language, num_results
                    )
                    if not competitors_data:
                        st.error("❌ Aucun concurrent trouvé")
                        return
                    st.session_state.competitors_data = competitors_data
                    st.session_state.search_completed = True
                    st.success(f"✅ {len(competitors_data)} concurrents trouvés et analysés")
                except Exception as e:
                    st.error(f"❌ Erreur: {e}")
                    return
        
        if st.session_state.search_completed and st.session_state.competitors_data:
            competitors_data = st.session_state.competitors_data
            
            with st.expander("📊 Données extraites", expanded=True):
                debug_data = [{
                    "Pos.": c.position,
                    "URL": c.url[:60] + "..." if len(c.url) > 60 else c.url,
                    "Title": c.title[:60] + "..." if len(c.title) > 60 else c.title,
                    "Hn": ", ".join(c.headings.keys()) if c.headings else "∅",
                    "Score": f"{c.relevance_score:.2f}",
                    "OK": "✅" if c.extraction_success else "❌"
                } for c in competitors_data]
                st.dataframe(pd.DataFrame(debug_data), use_container_width=True)
            
            with st.expander("📋 Structure Hn détaillée des concurrents"):
                headings_data = []
                for comp in competitors_data:
                    if comp.extraction_success and comp.headings:
                        for level, hlist in comp.headings.items():
                            for h in hlist:
                                headings_data.append({"Pos.": comp.position, "URL": urlparse(comp.url).netloc, "Level": level, "Heading": h})
                if headings_data:
                    st.dataframe(pd.DataFrame(headings_data).sort_values(['Pos.', 'Level']), use_container_width=True)
            
            # ÉTAPE 2 : SÉLECTION
            st.header("📋 Étape 2 : Sélection des concurrents")
            selected_indices = st.multiselect(
                "Choisissez les concurrents à inclure",
                range(len(competitors_data)),
                format_func=lambda x: f"#{competitors_data[x].position} - {competitors_data[x].title[:70]}... ({competitors_data[x].relevance_score:.2f})",
                default=list(range(min(5, len(competitors_data))))
            )
            if selected_indices:
                st.session_state.selected_competitors = [competitors_data[i] for i in selected_indices]
                st.success(f"✅ {len(selected_indices)} concurrents sélectionnés")
            else:
                st.warning("⚠️ Sélectionnez au moins un concurrent")
                return
            
            # ÉTAPE 3 : STRUCTURE Hn
            st.header("🏗️ Étape 3 : Génération de la structure Hn")
            structure_prompt = PromptTemplates.get_structure_prompt(target_keyword, st.session_state.selected_competitors)
            with st.expander("🔧 Voir/modifier le prompt de structure"):
                custom_structure_prompt = st.text_area("Prompt structure Hn", value=structure_prompt, height=300, key="structure_prompt_area")
            
            if st.button("🚀 Générer la structure Hn"):
                with st.spinner("Génération de la structure Hn..."):
                    structure_result = st.session_state.generator.ai_analyzer.analyze_with_custom_prompt(custom_structure_prompt, max_tokens=3000)
                    st.session_state.structure_result = structure_result
                    st.success("✅ Structure Hn générée")
            
            # ÉTAPE 4 : VALIDATION STRUCTURE
            if st.session_state.structure_result:
                st.header("✏️ Étape 4 : Validation de la structure Hn")
                st.code(st.session_state.structure_result, language=None)
                
                st.markdown("**✏️ Modifiez la structure ci-dessous :**")
                editable_structure = st.text_area(
                    "Structure Hn éditable",
                    value=st.session_state.structure_result,
                    height=400, key="editable_structure"
                )
                parsed_nodes = HeadingParser.parse_structure_text(editable_structure)
                
                if parsed_nodes:
                    with st.expander("👁️ Aperçu", expanded=True):
                        for node in parsed_nodes:
                            indent = "&nbsp;" * 6 * (node.level - 1)
                            weight = "**" if node.level <= 2 else ""
                            st.markdown(f"{indent}{weight}H{node.level}: {node.text}{weight}", unsafe_allow_html=True)
                    st.info(f"📊 {len(parsed_nodes)} headings : {sum(1 for n in parsed_nodes if n.level==1)} H1, {sum(1 for n in parsed_nodes if n.level==2)} H2, {sum(1 for n in parsed_nodes if n.level==3)} H3, {sum(1 for n in parsed_nodes if n.level==4)} H4")
                
                col_val, col_dl = st.columns([2, 1])
                with col_val:
                    if st.button("✅ Valider cette structure", type="primary"):
                        st.session_state.validated_structure = editable_structure
                        st.session_state.validated_nodes = parsed_nodes
                        st.success("✅ Structure validée !")
                
                with col_dl:
                    # Télécharger la structure Hn seule
                    if parsed_nodes:
                        hn_doc_buffer = WordGenerator.create_structure_only_document(
                            target_keyword, editable_structure,
                            title=st.session_state.get('title_result', ''),
                            meta_desc=st.session_state.get('meta_result', ''),
                            country=selected_country, language=selected_language
                        )
                        st.download_button(
                            label="📥 Télécharger la structure Hn (.docx)",
                            data=hn_doc_buffer,
                            file_name=f"structure_hn_{target_keyword.replace(' ', '_')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
            
            # ÉTAPE 5 : TITLE & META
            if st.session_state.get('validated_structure'):
                st.header("🏷️ Étape 5 : Title & Meta Description")
                col_t, col_m = st.columns(2)
                
                with col_t:
                    st.subheader("Title")
                    if st.button("🚀 Générer le Title"):
                        with st.spinner("Génération..."):
                            st.session_state.title_result = st.session_state.generator.ai_analyzer.analyze_with_custom_prompt(
                                PromptTemplates.get_title_prompt(target_keyword, st.session_state.selected_competitors), max_tokens=200
                            )
                    if st.session_state.title_result:
                        st.session_state.title_result = st.text_input("Title (modifiable)", value=st.session_state.title_result, key="title_edit")
                
                with col_m:
                    st.subheader("Meta Description")
                    if st.button("🚀 Générer la Meta Description"):
                        with st.spinner("Génération..."):
                            st.session_state.meta_result = st.session_state.generator.ai_analyzer.analyze_with_custom_prompt(
                                PromptTemplates.get_meta_description_prompt(target_keyword, st.session_state.selected_competitors), max_tokens=300
                            )
                    if st.session_state.meta_result:
                        st.session_state.meta_result = st.text_input("Meta Description (modifiable)", value=st.session_state.meta_result, key="meta_edit")
            
            # ÉTAPE 6 : RÉDACTION
            if (st.session_state.get('validated_nodes') and 
                st.session_state.get('title_result') and 
                st.session_state.get('meta_result')):
                
                st.header("✍️ Étape 6 : Rédaction de l'article")
                
                nodes = st.session_state.validated_nodes
                
                # Aperçu du prompt de rédaction (modifiable)
                with st.expander("🔧 Voir/modifier le prompt de rédaction (template)", expanded=False):
                    st.markdown("""
Le prompt ci-dessous est le **template** utilisé pour rédiger CHAQUE section. Les variables entre `{accolades}` seront remplacées automatiquement.

**Variables disponibles :** `{HEADING}`, `{HEADING_LEVEL}`, `{KEYWORD}`, `{PARENT_HEADING}`, `{FULL_STRUCTURE}`, `{PREVIOUS_SUMMARY}`, `{COMPETITOR_EXTRACTS}`, `{INTERNAL_LINKS}`, `{CLIENT_BRIEF}`, `{LANGUAGE}`

⚠️ **Laisser vide** pour utiliser le prompt par défaut (recommandé).
""")
                    custom_writing_template = st.text_area(
                        "Template de prompt de rédaction (laisser vide = prompt par défaut)",
                        value="",
                        height=300,
                        key="custom_writing_template",
                        placeholder="Laissez vide pour utiliser le prompt par défaut optimisé.\n\nOu écrivez votre propre template en utilisant les variables {HEADING}, {KEYWORD}, etc."
                    )
                
                # Aperçu d'un prompt généré (lecture seule)
                with st.expander("👁️ Aperçu d'un prompt de rédaction (premier H2)"):
                    first_h2 = next((n for n in nodes if n.level == 2), None)
                    if first_h2:
                        sample_prompt = PromptTemplates.get_paragraph_writing_prompt(
                            keyword=target_keyword,
                            current_heading=first_h2.text,
                            heading_level=2,
                            parent_heading="",
                            full_structure=HeadingParser.nodes_to_text(nodes),
                            previous_content_summary="",
                            competitor_extracts="[Extraits concurrents chargés dynamiquement]",
                            language=language_code,
                            internal_links_formatted=internal_links_formatted,
                            client_brief=full_client_brief
                        )
                        st.code(sample_prompt, language=None)
                
                if st.button("✍️ Lancer la rédaction complète", type="primary"):
                    progress_bar = st.progress(0)
                    status_text = st.empty()
                    
                    article_content = run_writing_engine(
                        nodes=nodes,
                        target_keyword=target_keyword,
                        selected_competitors=st.session_state.selected_competitors,
                        ai_analyzer=st.session_state.generator.ai_analyzer,
                        language_code=language_code,
                        internal_links_formatted=internal_links_formatted,
                        client_brief=full_client_brief,
                        writing_prompt_template=custom_writing_template,
                        progress_callback=lambda p: progress_bar.progress(p),
                        status_callback=lambda s: status_text.markdown(s)
                    )
                    
                    progress_bar.progress(1.0)
                    status_text.markdown("**✅ Rédaction terminée !**")
                    st.session_state.article_content = article_content
                    st.session_state.writing_done = True
                    st.success("✅ Article rédigé avec succès !")
            
            # AFFICHAGE & EXPORTS
            if st.session_state.writing_done and st.session_state.article_content:
                _display_article_and_exports(
                    st.session_state.article_content, target_keyword,
                    st.session_state.title_result, st.session_state.meta_result,
                    st.session_state.validated_structure, selected_country, selected_language,
                    st.session_state.selected_competitors, internal_links_input
                )
    
    # Reset
    if st.session_state.get('search_completed'):
        st.divider()
        if st.button("🔄 Nouvelle recherche"):
            for key in list(st.session_state.keys()):
                if key != 'generator':
                    del st.session_state[key]
            st.rerun()


def _display_article_and_exports(article_content, target_keyword, title_result, meta_result,
                                  validated_structure, selected_country, selected_language,
                                  selected_competitors, internal_links_input):
    """Affiche l'article et les boutons d'export"""
    st.header("📄 Article rédigé")

    # Construire le contenu markdown complet pour l'affichage ET la copie
    full_markdown = ""
    for node in article_content:
        level_map = {1: "#", 2: "##", 3: "###", 4: "####"}
        prefix = level_map.get(node.level, "####")
        full_markdown += f"{prefix} {node.text}\n\n"
        if node.content:
            full_markdown += f"{node.content}\n\n"

    st.markdown(full_markdown)

    total_words = sum(len(n.content.split()) for n in article_content if n.content)
    st.info(f"📊 **Statistiques** : {total_words} mots | {len(article_content)} sections")

    # Bouton copier le contenu
    with st.expander("📋 Copier l'article (Markdown)"):
        st.caption("Cliquez sur l'icône 📋 en haut à droite du bloc pour copier tout le contenu.")
        st.code(full_markdown, language="markdown")

    st.subheader("📥 Exports")
    col_e1, col_e2, col_e3 = st.columns(3)
    
    with col_e1:
        try:
            article = ContentArticle(
                target_keyword=target_keyword, title=title_result,
                meta_description=meta_result, headings_tree=article_content,
                country=selected_country, language=selected_language,
                competitors_analyzed=[c.url for c in selected_competitors],
                internal_links_context=internal_links_input or ""
            )
            word_buffer = WordGenerator.create_full_article_document(article)
            st.download_button(
                label="📄 Article complet (.docx)",
                data=word_buffer,
                file_name=f"article_seo_{target_keyword.replace(' ', '_')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        except Exception as e:
            st.error(f"Erreur: {e}")
    
    with col_e2:
        try:
            seo_brief = SEOBrief(
                target_keyword=target_keyword, optimized_title=title_result,
                optimized_meta_description=meta_result, headings_structure=validated_structure,
                country=selected_country, language=selected_language,
                competitors_analyzed=[c.url for c in selected_competitors]
            )
            brief_buffer = WordGenerator.create_seo_brief_document(seo_brief)
            st.download_button(
                label="📋 Brief SEO (.docx)",
                data=brief_buffer,
                file_name=f"brief_seo_{target_keyword.replace(' ', '_')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        except Exception as e:
            st.error(f"Erreur: {e}")
    
    with col_e3:
        try:
            hn_buffer = WordGenerator.create_structure_only_document(
                target_keyword, validated_structure,
                title=title_result, meta_desc=meta_result,
                country=selected_country, language=selected_language
            )
            st.download_button(
                label="📐 Structure Hn (.docx)",
                data=hn_buffer,
                file_name=f"structure_hn_{target_keyword.replace(' ', '_')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        except Exception as e:
            st.error(f"Erreur: {e}")


main()