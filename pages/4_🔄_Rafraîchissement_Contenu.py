"""
🔄 RAFRAÎCHISSEMENT DE CONTENU — Gérer Seul
=============================================
Outil de rafraîchissement d'articles existants :
- Scraping de l'article existant + structures concurrentes
- Analyse SERP via DataForSEO
- Structure optimale + rédaction par Claude
- Maillage interne via export Screaming Frog (URL + H1 + Title)
- Mode bulk (jusqu'à 5 articles)
- Export .docx
"""

import streamlit as st
import requests
from bs4 import BeautifulSoup, Comment
import base64
import anthropic
import pandas as pd
import json
import re
import io
import zipfile
import time
from datetime import date, timedelta
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from urllib.parse import urlparse

from utils.auth import check_password

# =============================================================================
# CONFIGURATION
# =============================================================================
st.set_page_config(
    page_title="Rafraîchissement Contenu | Ma Toolbox SEO",
    page_icon="🔄",
    layout="wide",
    initial_sidebar_state="expanded"
)

check_password()

DATE_CONTEXT = "Nous sommes en 2026. La date du jour est le 27 mars 2026."

EDITORIAL_GUIDELINES = f"""Tu rédiges pour le site "Gérer Seul", expert reconnu de l'immobilier.
- Ton : professionnel, expert, autoritaire mais accessible
- Expertise : immobilier, juridique, lois, réglementations, gestion locative, copropriété
- Tu es incolable sur tous les sujets juridiques et de loi liés à l'immobilier
- Vouvoiement
- Précision juridique : cite les articles de loi, décrets, dates d'entrée en vigueur quand pertinent

⚠️⚠️⚠️ RÈGLE ABSOLUE SUR LES DATES ⚠️⚠️⚠️
{DATE_CONTEXT}
- TOUTES les informations (lois, chiffres, montants, seuils, barèmes, taux, réglementations) DOIVENT être celles en vigueur en 2026.
- Si tu mentionnes une date, un montant, un barème ou un seuil, il DOIT correspondre à la réalité de 2026.
- NE CITE JAMAIS "2024" ou "2025" comme année en cours. L'année en cours est 2026.
- Si l'article existant contient des dates 2024 ou 2025, tu DOIS les mettre à jour pour 2026.
- En cas de doute sur un chiffre précis 2026, indique "en 2026" et donne la tendance plutôt qu'un chiffre obsolète de 2024/2025.
⚠️⚠️⚠️ FIN RÈGLE DATES ⚠️⚠️⚠️"""


# =============================================================================
# SCRAPING
# =============================================================================
class ContentScraper:
    """Scrape des articles et structures Hn."""

    HEADERS = {
        "User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
                      "AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
    }

    @staticmethod
    def scrape_article(url: str) -> dict:
        """Scrape le contenu complet d'un article."""
        try:
            resp = requests.get(url, headers=ContentScraper.HEADERS, timeout=15)
            resp.raise_for_status()
            soup = BeautifulSoup(resp.text, "lxml")

            # Nettoyer les éléments non-contenu
            for tag in soup.find_all(["script", "style", "nav", "footer", "aside", "header"]):
                tag.decompose()
            for comment in soup.find_all(string=lambda t: isinstance(t, Comment)):
                comment.extract()

            # Extraire title et meta description
            title = ""
            title_tag = soup.find("title")
            if title_tag:
                title = title_tag.get_text(strip=True)

            meta_desc = ""
            meta_tag = soup.find("meta", attrs={"name": "description"})
            if meta_tag:
                meta_desc = meta_tag.get("content", "")

            # Extraire H1
            h1 = ""
            h1_tag = soup.find("h1")
            if h1_tag:
                h1 = h1_tag.get_text(strip=True)

            # Trouver la zone de contenu principale
            content_area = (
                soup.find("article")
                or soup.find("main")
                or soup.find("div", class_=re.compile(r'(content|article|post|entry)', re.I))
                or soup.find("body")
            )

            # Extraire le contenu structuré
            # Extraire les liens internes existants
            existing_links = []
            target_domain = urlparse(url).netloc
            if content_area:
                seen_urls = set()
                for a_tag in content_area.find_all("a", href=True):
                    href = a_tag.get("href", "").strip()
                    anchor = a_tag.get_text(strip=True)
                    if not href or not anchor or len(anchor) < 3:
                        continue
                    if href.startswith("/"):
                        href = f"https://{target_domain}{href}"
                    if target_domain in href and href not in seen_urls:
                        seen_urls.add(href)
                        existing_links.append({"url": href, "anchor": anchor})

            content_parts = []
            hn_structure = []
            if content_area:
                for elem in content_area.find_all(
                    ["h1", "h2", "h3", "h4", "h5", "h6", "p", "ul", "ol", "li", "blockquote"]
                ):
                    tag_name = elem.name
                    text = elem.get_text(strip=True)
                    if not text:
                        continue
                    if tag_name.startswith("h"):
                        level = int(tag_name[1])
                        content_parts.append(f"{'#' * level} {text}")
                        hn_structure.append({"level": level, "text": text})
                    elif tag_name == "li":
                        content_parts.append(f"- {text}")
                    else:
                        content_parts.append(text)

            content_text = "\n\n".join(content_parts)

            return {
                "url": url,
                "title": title,
                "meta_description": meta_desc,
                "h1": h1,
                "content": content_text,
                "hn_structure": hn_structure,
                "existing_links": existing_links,
                "word_count": len(content_text.split()),
                "success": True,
            }
        except Exception as e:
            return {"url": url, "success": False, "error": str(e), "word_count": 0}

    @staticmethod
    def scrape_hn_structure(url: str) -> dict:
        """Extrait uniquement la structure Hn d'une page."""
        try:
            resp = requests.get(url, headers=ContentScraper.HEADERS, timeout=10)
            resp.raise_for_status()
            soup = BeautifulSoup(resp.text, "lxml")

            title = ""
            title_tag = soup.find("title")
            if title_tag:
                title = title_tag.get_text(strip=True)

            hn_structure = []
            for tag in soup.find_all(re.compile(r'^h[1-6]$')):
                level = int(tag.name[1])
                text = tag.get_text(strip=True)
                if text and len(text) > 2:
                    hn_structure.append({"level": level, "text": text})

            return {"url": url, "title": title, "structure": hn_structure, "success": True}
        except Exception:
            return {"url": url, "structure": [], "success": False}


# =============================================================================
# DATAFORSEO
# =============================================================================
class DataForSEOClient:
    """Client DataForSEO pour les SERP."""

    def __init__(self, username: str, password: str):
        self.base_url = "https://api.dataforseo.com/v3"
        self.session = requests.Session()
        creds = base64.b64encode(f"{username}:{password}".encode()).decode()
        self.session.headers.update({
            "Authorization": f"Basic {creds}",
            "Content-Type": "application/json",
        })
        self.total_cost = 0.0

    def get_serp(self, keyword: str, location_code: int = 2250,
                 language_code: str = "fr", limit: int = 10) -> list:
        """Récupère les résultats SERP organiques pour un mot-clé."""
        endpoint = f"{self.base_url}/serp/google/organic/live/advanced"
        payload = [{
            "keyword": keyword,
            "location_code": location_code,
            "language_code": language_code,
            "device": "desktop",
            "os": "windows",
            "depth": limit,
        }]

        resp = self.session.post(endpoint, json=payload)
        resp.raise_for_status()
        data = resp.json()
        self.total_cost += data.get("cost", 0)

        results = []
        tasks = data.get("tasks", [])
        if tasks and tasks[0].get("status_code") == 20000:
            items = tasks[0].get("result", [{}])[0].get("items", [])
            for item in items:
                if item.get("type") == "organic":
                    results.append({
                        "position": item.get("rank_group", 0),
                        "url": item.get("url", ""),
                        "title": item.get("title", ""),
                        "domain": item.get("domain", ""),
                    })
        return results


# =============================================================================
# MAILLAGE INTERNE
# =============================================================================
def parse_screaming_frog_csv(uploaded_file) -> pd.DataFrame:
    """Parse un export CSV Screaming Frog (colonnes: Address, Title 1, H1-1)."""
    content = uploaded_file.getvalue()

    df = None
    for encoding in ["utf-8", "utf-8-sig", "latin-1"]:
        for sep in [",", ";", "\t"]:
            try:
                candidate = pd.read_csv(io.BytesIO(content), encoding=encoding, sep=sep)
                if len(candidate.columns) > 1:
                    df = candidate
                    break
            except Exception:
                continue
        if df is not None:
            break

    if df is None:
        raise ValueError("Impossible de lire le fichier CSV.")

    col_mapping = {}
    for col in df.columns:
        cl = col.lower().strip()
        if cl in ("address", "url", "adresse"):
            col_mapping[col] = "url"
        elif cl in ("title 1", "title", "titre 1", "titre"):
            col_mapping[col] = "title"
        elif cl in ("h1-1", "h1", "h1 1"):
            col_mapping[col] = "h1"

    df = df.rename(columns=col_mapping)

    available = [c for c in ["url", "title", "h1"] if c in df.columns]
    if "url" not in available:
        raise ValueError(
            "Colonne URL introuvable (attendu : 'Address'). "
            f"Colonnes trouvées : {', '.join(df.columns.tolist()[:10])}"
        )

    df = df[available].dropna(subset=["url"])
    df = df[df["url"].str.startswith("http")]
    return df


def _build_gsc_service():
    """Construit le service GSC à partir des secrets Streamlit."""
    try:
        from googleapiclient.discovery import build
        if "GSC_SERVICE_ACCOUNT" in st.secrets:
            from google.oauth2 import service_account
            sa_info = dict(st.secrets["GSC_SERVICE_ACCOUNT"])
            creds = service_account.Credentials.from_service_account_info(
                sa_info,
                scopes=["https://www.googleapis.com/auth/webmasters.readonly"],
            )
            return build("searchconsole", "v1", credentials=creds)
        from google.oauth2.credentials import Credentials
        client_id = st.secrets.get("GSC_CLIENT_ID", "")
        client_secret = st.secrets.get("GSC_CLIENT_SECRET", "")
        refresh_token = st.secrets.get("GSC_REFRESH_TOKEN", "")
        if client_id and client_secret and refresh_token:
            creds = Credentials(
                token=None,
                refresh_token=refresh_token,
                client_id=client_id,
                client_secret=client_secret,
                token_uri="https://oauth2.googleapis.com/token",
            )
            return build("searchconsole", "v1", credentials=creds)
    except Exception:
        pass
    return None


@st.cache_data(ttl=3600)
def fetch_top_gsc_urls(site_url: str = "sc-domain:gererseul.com", limit: int = 500) -> list:
    """Récupère les top URLs par clics (6 derniers mois). Mis en cache 1h."""
    service = _build_gsc_service()
    if not service:
        return []
    end_date = date.today() - timedelta(days=3)
    start_date = end_date - timedelta(days=180)
    body = {
        "startDate": start_date.strftime("%Y-%m-%d"),
        "endDate": end_date.strftime("%Y-%m-%d"),
        "dimensions": ["page"],
        "rowLimit": limit,
        "dataState": "final",
    }
    try:
        resp = service.searchanalytics().query(siteUrl=site_url, body=body).execute()
        rows = resp.get("rows", [])
        return [
            {"url": row["keys"][0], "topic": url_to_topic(row["keys"][0])}
            for row in rows
            if row.get("clicks", 0) > 0
        ]
    except Exception:
        return []


def url_to_topic(url: str) -> str:
    """Extrait le sujet depuis le slug d'URL."""
    path = urlparse(url).path
    parts = [p for p in path.strip("/").split("/") if p and len(p) > 2]
    if parts:
        return parts[-1].replace("-", " ").replace("_", " ")
    return ""


def build_linking_context(existing_links: list, candidate_pages: list, current_url: str) -> str:
    """Construit le contexte de maillage interne pour le prompt Claude."""
    parts = []

    if existing_links:
        parts.append("### Liens internes DÉJÀ PRÉSENTS dans l'article original (à réintégrer OBLIGATOIREMENT) :")
        for link in existing_links[:6]:
            parts.append(f"- URL : {link['url']} | Ancre d'origine : \"{link['anchor']}\"")
        parts.append("")

    pages = [p for p in candidate_pages if p["url"] != current_url][:150]
    if pages:
        parts.append("### Pages du site disponibles pour le maillage (top clics GSC) :")
        for page in pages:
            parts.append(f"- {page['url']}  →  {page['topic']}")
        parts.append("")

    if not parts:
        return ""

    context = "\n".join(parts)
    rules = """
CONSIGNES MAILLAGE INTERNE :

**Règle 1 — Liens existants (priorité absolue)**
Si des liens existants sont listés ci-dessus, tu DOIS les réintégrer dans le nouvel article avec la même URL. Adapte l'ancre à la phrase dans laquelle tu les places (l'ancre peut changer, l'URL reste identique).

**Règle 2 — Nouveaux liens (pour compléter jusqu'à 6 au total)**
Parmi les pages disponibles, sélectionne celles dont la thématique est DIRECTEMENT connexe au sujet de l'article. Pour chaque lien retenu :
- Utilise une ancre optimisée : le mot-clé principal de la page cible (déduit de son slug ou topic), intégré naturellement dans une phrase.
- Si l'ancre optimisée ne s'intègre pas naturellement, utilise une ancre descriptive acceptable.

**Règle 3 — Format et contraintes**
- Maximum 6 liens dans tout l'article (liens existants + nouveaux)
- Format : [texte ancre](url_complète)
- Liens dans le corps de l'article uniquement, répartis sur l'ensemble de l'article
- Si aucune page n'est thématiquement connexe, ne force pas de lien
"""
    return context + rules


# =============================================================================
# CLAUDE — STRUCTURE & RÉDACTION
# =============================================================================
def generate_structure(api_key: str, keyword: str, existing_content: dict,
                       competitor_structures: list) -> dict:
    """Phase 1 : Claude analyse et propose la structure optimale."""
    client = anthropic.Anthropic(api_key=api_key)

    # Formater la structure existante
    existing_hn = ""
    if existing_content.get("hn_structure"):
        existing_hn = "\n".join([
            f"{'  ' * (h['level'] - 1)}H{h['level']}: {h['text']}"
            for h in existing_content["hn_structure"]
        ])

    # Formater les structures concurrentes
    comp_text = ""
    for i, comp in enumerate(competitor_structures):
        if comp.get("success") and comp.get("structure"):
            comp_text += f"\n### Concurrent {i + 1} : {comp['url']}\n"
            for h in comp["structure"][:30]:
                comp_text += f"{'  ' * (h['level'] - 1)}H{h['level']}: {h['text']}\n"

    prompt = f"""CONTEXTE TEMPOREL : {DATE_CONTEXT} Toute information doit être à jour pour 2026. Ne mentionne JAMAIS 2024 ou 2025 comme année en cours.

Tu es un expert SEO spécialisé en immobilier et juridique (site Gérer Seul).

## Article existant à rafraîchir :
- URL : {existing_content.get('url', 'N/A')}
- Title actuel : {existing_content.get('title', 'N/A')}
- Meta description actuelle : {existing_content.get('meta_description', 'N/A')}
- H1 actuel : {existing_content.get('h1', 'N/A')}
- Nombre de mots : {existing_content.get('word_count', 'N/A')}

### Structure Hn actuelle :
{existing_hn or 'Non disponible'}

## Structures des 5 meilleurs résultats Google sur "{keyword}" :
{comp_text or 'Aucune structure récupérée'}

## Ta mission :
1. Identifie les sujets ESSENTIELS que les concurrents couvrent et que l'article actuel ne couvre pas
2. Élimine les redondances : si 3 concurrents ont des sections qui disent la même chose, n'en garde qu'une
3. Propose une structure Hn OPTIMALE :
   - Couvre toute la sémantique nécessaire pour "{keyword}"
   - Chaque section apporte une valeur UNIQUE (pas de doublons)
   - Structure digeste et logique (pas de structure à rallonge)
   - Meilleure que l'existant ET meilleure que chaque concurrent individuellement

4. RÈGLE ABSOLUE : N'inclus JAMAIS de section "modèle", "template", "exemple de lettre/courrier" dans la structure. Gérer Seul ne propose pas de modèles sur son blog. Si les concurrents en ont, ignore ces sections.

5. Propose :
   - Le meilleur title SEO (< 60 caractères, mot-clé en début si possible)
   - La meilleure meta description (< 155 caractères, incitative, avec mot-clé)
   - Le H1 optimal

## FORMAT JSON OBLIGATOIRE :
```json
{{
  "title_seo": "...",
  "meta_description": "...",
  "h1": "...",
  "structure": [
    {{"level": 2, "text": "Titre H2"}},
    {{"level": 3, "text": "Sous-titre H3"}},
    {{"level": 2, "text": "Autre H2"}}
  ],
  "rationale": "Explication courte de tes choix"
}}
```"""

    response = client.messages.create(
        model="claude-sonnet-4-20250514",
        max_tokens=3000,
        system=f"Tu es un rédacteur SEO expert en immobilier pour le site Gérer Seul. {DATE_CONTEXT} Toute information (lois, montants, barèmes, seuils, taux) doit être celle en vigueur en 2026. Ne mentionne JAMAIS 2024 ou 2025 comme année en cours.",
        messages=[{"role": "user", "content": prompt}],
    )

    text = response.content[0].text
    json_match = re.search(r'```json\s*(.*?)\s*```', text, re.DOTALL)
    if json_match:
        return json.loads(json_match.group(1))
    json_match = re.search(r'\{[\s\S]*\}', text)
    if json_match:
        return json.loads(json_match.group(0))
    return {}


def write_article(api_key: str, keyword: str, structure: dict,
                  existing_content: dict, linking_data: str) -> str:
    """Phase 2 : Claude rédige l'article complet rafraîchi."""
    client = anthropic.Anthropic(api_key=api_key)

    structure_text = "\n".join([
        f"{'#' * h['level']} {h['text']}"
        for h in structure.get("structure", [])
    ])

    existing_text = existing_content.get("content", "Non disponible")
    # Tronquer si trop long mais garder un maximum
    if len(existing_text) > 12000:
        existing_text = existing_text[:12000] + "\n\n[...contenu tronqué...]"

    linking_section = ""
    if linking_data:
        linking_section = f"\n## MAILLAGE INTERNE\n{linking_data[:7000]}"

    prompt = f"""{EDITORIAL_GUIDELINES}

## MISSION : Rédige l'article COMPLET rafraîchi pour le mot-clé "{keyword}"

### Balises :
- Title SEO : {structure.get('title_seo', '')}
- Meta description : {structure.get('meta_description', '')}
- H1 : {structure.get('h1', '')}

### Structure Hn à suivre OBLIGATOIREMENT :
{structure_text}

### Contenu existant de l'article (à AMÉLIORER, pas à copier) :
{existing_text}

{linking_section}

## CONSIGNES DE RÉDACTION :

1. **Contenu** :
   - Suis la structure Hn ci-dessus EXACTEMENT (chaque H2/H3 doit être présent)
   - Améliore le contenu : infos à jour, couverture sémantique complète
   - Cite les articles de loi, décrets, dates quand pertinent
   - Chaque section doit apporter une vraie valeur au lecteur
   - Ne répète pas les mêmes informations d'une section à l'autre
   - N'inclus JAMAIS de modèle, template, exemple de lettre ou de courrier. Gérer Seul ne propose pas de modèles.

2. **Format de sortie** :
   - Markdown avec ## pour H2, ### pour H3, etc.
   - Liens internes en format [ancre](url)
   - NE COMMENCE PAS par le H1 (il est déjà défini séparément)
   - Commence directement par le contenu du premier H2
   - Pas de texte en gras, pas de mise en forme spéciale

RAPPEL CRITIQUE : {DATE_CONTEXT} Chaque date, montant, seuil, barème ou loi mentionné DOIT être la version 2026. Si l'article original dit "en 2024" ou "en 2025", remplace par la version 2026 en vigueur. Ne laisse AUCUNE référence à 2024 ou 2025 comme année en cours dans l'article.

Rédige l'article complet maintenant. Sois exhaustif sur chaque section."""

    response = client.messages.create(
        model="claude-sonnet-4-20250514",
        max_tokens=8000,
        system=f"Tu es un rédacteur SEO expert en immobilier pour le site Gérer Seul. {DATE_CONTEXT} RÈGLE ABSOLUE : toute information (lois, montants, barèmes, seuils, taux, plafonds) doit être celle en vigueur en 2026. Si tu ne connais pas le chiffre exact 2026, écris 'en 2026' et donne la tendance. Ne cite JAMAIS 2024 ou 2025 comme année en cours. L'article existant peut contenir des dates obsolètes — remplace-les TOUTES par 2026.",
        messages=[{"role": "user", "content": prompt}],
    )

    return response.content[0].text


# =============================================================================
# GÉNÉRATION DOCX
# =============================================================================
def add_hyperlink(paragraph, text, url):
    """Ajoute un hyperlien cliquable dans un paragraphe docx."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    # Soulignement simple pour identifier les liens
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)

    run.append(rPr)
    run_text = OxmlElement("w:t")
    run_text.set(qn("xml:space"), "preserve")
    run_text.text = text
    run.append(run_text)
    hyperlink.append(run)
    paragraph._element.append(hyperlink)


def _add_text_with_links(paragraph, text: str):
    """Ajoute du texte avec des liens [ancre](url) convertis en hyperlinks docx."""
    # Nettoyer le markdown bold/italic
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'(?<!\*)\*(?!\*)(.+?)\*', r'\1', text)

    pattern = r'\[([^\]]+)\]\(([^)]+)\)'
    last_end = 0

    for match in re.finditer(pattern, text):
        before = text[last_end:match.start()]
        if before:
            run = paragraph.add_run(before)
            run.font.size = Pt(11)

        anchor_text = match.group(1)
        link_url = match.group(2)
        add_hyperlink(paragraph, anchor_text, link_url)

        last_end = match.end()

    remaining = text[last_end:]
    if remaining:
        run = paragraph.add_run(remaining)
        run.font.size = Pt(11)


def add_black_heading(doc, text, level):
    """Ajoute un heading avec la couleur forcée en noir."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return heading


def create_docx(url: str, title_seo: str, meta_desc: str, h1: str,
                article_markdown: str) -> io.BytesIO:
    """Génère un fichier .docx avec le format demandé."""
    doc = Document()

    # Style de base
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    # Forcer tous les styles Heading en noir
    for i in range(1, 7):
        style_name = f"Heading {i}"
        if style_name in doc.styles:
            doc.styles[style_name].font.color.rgb = RGBColor(0, 0, 0)

    # Ligne 1 : URL
    doc.add_paragraph(url)

    # Ligne 2 : Title SEO
    doc.add_paragraph(f"title SEO : {title_seo}")

    # Ligne 3 : Meta description
    doc.add_paragraph(f"meta description : {meta_desc}")

    # Ligne 4 : H1
    add_black_heading(doc, h1, level=1)

    # Corps de l'article
    lines = article_markdown.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].strip()

        if not line:
            i += 1
            continue

        # Détection des titres Hn
        heading_match = re.match(r'^(#{2,6})\s+(.+)$', line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            # Nettoyer le markdown
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            add_black_heading(doc, text, level=level)
            i += 1
            continue

        # Éléments de liste
        list_match = re.match(r'^[-*]\s+(.+)$', line)
        numbered_match = re.match(r'^\d+\.\s+(.+)$', line)
        if list_match or numbered_match:
            text = list_match.group(1) if list_match else numbered_match.group(1)
            style_name = "List Bullet" if list_match else "List Number"
            p = doc.add_paragraph(style=style_name)
            _add_text_with_links(p, text)
            i += 1
            continue

        # Paragraphe normal — regrouper les lignes consécutives
        para_lines = [line]
        i += 1
        while i < len(lines):
            next_line = lines[i].strip()
            if (not next_line
                    or re.match(r'^#{2,6}\s+', next_line)
                    or next_line.startswith("- ")
                    or next_line.startswith("* ")
                    or re.match(r'^\d+\.\s', next_line)):
                break
            para_lines.append(next_line)
            i += 1

        full_text = " ".join(para_lines)
        p = doc.add_paragraph()
        _add_text_with_links(p, full_text)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# =============================================================================
# TRAITEMENT D'UN ARTICLE
# =============================================================================
def process_article(url: str, keyword: str, api: DataForSEOClient,
                    anthropic_key: str, sf_data=None,
                    progress_callback=None) -> dict:
    """Pipeline complet pour rafraîchir un article."""

    def update(msg):
        if progress_callback:
            progress_callback(msg)

    # 1. Scraper l'article existant
    update("🔍 Scraping de l'article existant...")
    existing = ContentScraper.scrape_article(url)
    if not existing.get("success"):
        return {
            "success": False,
            "error": f"Impossible de scraper {url} : {existing.get('error', 'Erreur inconnue')}",
        }
    update(f"  ✅ Article récupéré ({existing['word_count']} mots)")

    # 2. SERP
    update(f"🔍 Analyse SERP pour \"{keyword}\"...")
    try:
        serp_results = api.get_serp(keyword)
    except requests.exceptions.HTTPError as e:
        return {"success": False, "error": f"Erreur DataForSEO : {e}"}
    update(f"  ✅ {len(serp_results)} résultats SERP")

    # 3. Scraper les structures concurrentes (top 5 hors même domaine)
    update("🔍 Scraping des structures concurrentes...")
    target_domain = urlparse(url).netloc
    competitor_structures = []
    for result in serp_results[:7]:
        if urlparse(result["url"]).netloc == target_domain:
            continue
        if len(competitor_structures) >= 5:
            break
        structure = ContentScraper.scrape_hn_structure(result["url"])
        competitor_structures.append(structure)
        time.sleep(0.3)

    successful = sum(1 for c in competitor_structures if c.get("success"))
    update(f"  ✅ {successful}/{len(competitor_structures)} structures récupérées")

    # 4. Claude Phase 1 : Structure optimale
    update("🧠 Claude analyse et propose la structure optimale...")
    try:
        structure = generate_structure(
            api_key=anthropic_key,
            keyword=keyword,
            existing_content=existing,
            competitor_structures=competitor_structures,
        )
    except Exception as e:
        return {"success": False, "error": f"Erreur Claude (structure) : {e}"}

    if not structure or not structure.get("structure"):
        return {"success": False, "error": "Claude n'a pas pu générer la structure."}

    update(f"  ✅ Structure : {len(structure['structure'])} sections")

    # 5. Claude Phase 2 : Rédaction
    update("✍️ Claude rédige l'article rafraîchi...")

    # Construire le contexte de maillage interne
    existing_links = existing.get("existing_links", [])
    candidate_pages = []
    if sf_data is not None and not sf_data.empty:
        # CSV override
        for _, row in sf_data[sf_data["url"] != url].head(150).iterrows():
            topic = row.get("h1") or row.get("title") or url_to_topic(str(row.get("url", "")))
            candidate_pages.append({"url": str(row["url"]), "topic": str(topic)})
        update(f"🔗 Maillage : {len(candidate_pages)} pages depuis le CSV")
    else:
        # GSC (source principale)
        update("🔗 Récupération des top URLs depuis la GSC...")
        candidate_pages = fetch_top_gsc_urls()
        if candidate_pages:
            update(f"  ✅ {len(candidate_pages)} pages GSC pour le maillage")
        else:
            update("  ⚠️ GSC non accessible, maillage sur liens existants uniquement")

    linking_data = build_linking_context(existing_links, candidate_pages, url)

    try:
        article_content = write_article(
            api_key=anthropic_key,
            keyword=keyword,
            structure=structure,
            existing_content=existing,
            linking_data=linking_data,
        )
    except Exception as e:
        return {"success": False, "error": f"Erreur Claude (rédaction) : {e}"}

    update("  ✅ Article rédigé")

    # 6. Générer le DOCX
    update("📄 Génération du .docx...")
    docx_buffer = create_docx(
        url=url,
        title_seo=structure.get("title_seo", ""),
        meta_desc=structure.get("meta_description", ""),
        h1=structure.get("h1", ""),
        article_markdown=article_content,
    )
    update("  ✅ Document généré")

    return {
        "success": True,
        "url": url,
        "keyword": keyword,
        "title_seo": structure.get("title_seo", ""),
        "meta_desc": structure.get("meta_description", ""),
        "h1": structure.get("h1", ""),
        "structure": structure,
        "article_preview": article_content[:2000],
        "docx": docx_buffer,
        "existing_word_count": existing["word_count"],
        "new_word_count": len(article_content.split()),
    }


# =============================================================================
# INTERFACE
# =============================================================================
st.title("🔄 Rafraîchissement de Contenu")
st.markdown("*Outil dédié Gérer Seul — Rafraîchis tes articles pour regagner des positions.*")

# ─── Sidebar ───
with st.sidebar:
    st.header("⚙️ Configuration")
    dataforseo_username = st.text_input(
        "Username DataForSEO",
        value=st.secrets.get("DATAFORSEO_USERNAME", ""),
        type="password",
    )
    dataforseo_password = st.text_input(
        "Password DataForSEO",
        value=st.secrets.get("DATAFORSEO_PASSWORD", ""),
        type="password",
    )
    anthropic_key = st.text_input(
        "Clé API Claude",
        value=st.secrets.get("ANTHROPIC_API_KEY", ""),
        type="password",
    )

    st.divider()
    st.header("🔗 Maillage interne")
    st.markdown(
        "Le maillage utilise automatiquement les **top 500 URLs par clics** (GSC, 6 derniers mois).\n\n"
        "Upload un CSV Screaming Frog pour remplacer la GSC :"
    )

    sf_file = st.file_uploader("Export Screaming Frog (.csv) — optionnel", type=["csv"])

    if sf_file:
        try:
            sf_data = parse_screaming_frog_csv(sf_file)
            st.session_state["sf_data"] = sf_data
            st.success(f"✅ {len(sf_data)} pages CSV chargées (prioritaire sur la GSC)")
            with st.expander("Aperçu"):
                st.dataframe(sf_data.head(15), use_container_width=True)
        except Exception as e:
            st.error(f"Erreur : {e}")

    if "sf_data" in st.session_state:
        st.caption(f"📄 {len(st.session_state['sf_data'])} pages CSV en mémoire")
    else:
        with st.spinner("Connexion GSC..."):
            gsc_pages = fetch_top_gsc_urls()
        if gsc_pages:
            st.success(f"✅ {len(gsc_pages)} pages GSC chargées")
        else:
            st.warning("⚠️ GSC non accessible")

    st.divider()
    st.markdown("""
    ### Pipeline
    1. Scrape article existant
    2. SERP DataForSEO (top 10)
    3. Scrape structures concurrentes
    4. Claude → structure optimale
    5. Claude → article rafraîchi
    6. Export .docx
    """)

# ─── Validations ───
if not dataforseo_username or not dataforseo_password:
    st.warning("Configure tes identifiants DataForSEO dans la sidebar.")
    st.stop()

if not anthropic_key:
    st.warning("Configure ta clé API Claude dans la sidebar.")
    st.stop()

# ─── Tabs ───
tab_single, tab_bulk = st.tabs(["📝 Article unique", "📦 Mode bulk (5 articles)"])


# ─── Tab : Article unique ───
with tab_single:
    col1, col2 = st.columns([2, 1])
    with col1:
        single_url = st.text_input(
            "URL de l'article à rafraîchir",
            placeholder="https://www.gererseul.com/mon-article",
            key="single_url",
        )
    with col2:
        single_keyword = st.text_input(
            "Mot-clé cible",
            placeholder="bail commercial durée",
            key="single_kw",
        )

    if st.button("🚀 Rafraîchir l'article", type="primary",
                  use_container_width=True, key="btn_single"):
        if not single_url or not single_keyword:
            st.error("Renseigne l'URL et le mot-clé.")
            st.stop()

        api = DataForSEOClient(dataforseo_username, dataforseo_password)
        sf_data = st.session_state.get("sf_data")

        with st.status("Traitement en cours...", expanded=True) as status:
            result = process_article(
                url=single_url,
                keyword=single_keyword,
                api=api,
                anthropic_key=anthropic_key,
                sf_data=sf_data,
                progress_callback=lambda msg: st.write(msg),
            )
            if result["success"]:
                status.update(label="Article rafraîchi !", state="complete")
            else:
                status.update(label="Erreur", state="error")

        if result["success"]:
            st.divider()
            st.success(f"**{result['title_seo']}**")

            c1, c2, c3 = st.columns(3)
            c1.metric("Mots (avant)", result["existing_word_count"])
            c2.metric("Mots (après)", result["new_word_count"])
            c3.metric("Sections Hn", len(result["structure"].get("structure", [])))

            if result["structure"].get("rationale"):
                with st.expander("🧠 Pourquoi cette structure ?"):
                    st.markdown(result["structure"]["rationale"])

            with st.expander("📄 Aperçu de l'article", expanded=True):
                st.markdown(result["article_preview"] + "\n\n*[...]*")

            filename = f"{single_keyword}.docx"

            st.download_button(
                f"📥 Télécharger {filename}",
                data=result["docx"],
                file_name=filename,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )

            st.caption(f"Coût DataForSEO : ${api.total_cost:.4f}")
        else:
            st.error(result["error"])


# ─── Tab : Bulk ───
with tab_bulk:
    st.subheader("Rafraîchir jusqu'à 5 articles")

    bulk_items = []
    for i in range(5):
        col1, col2 = st.columns([2, 1])
        with col1:
            b_url = st.text_input(
                f"URL {i + 1}", key=f"bulk_url_{i}",
                placeholder="https://www.gererseul.com/...",
            )
        with col2:
            b_kw = st.text_input(
                f"Mot-clé {i + 1}", key=f"bulk_kw_{i}",
                placeholder="mot clé cible",
            )
        bulk_items.append({"url": b_url.strip(), "keyword": b_kw.strip()})

    valid_items = [item for item in bulk_items if item["url"] and item["keyword"]]

    if valid_items:
        st.caption(f"{len(valid_items)} article(s) à traiter")

    if st.button(
        f"🚀 Lancer le batch ({len(valid_items)} articles)",
        type="primary",
        use_container_width=True,
        key="btn_bulk",
        disabled=len(valid_items) == 0,
    ):
        api = DataForSEOClient(dataforseo_username, dataforseo_password)
        sf_data = st.session_state.get("sf_data")

        results = []
        progress = st.progress(0, text="Démarrage...")
        log_container = st.container()

        for idx, item in enumerate(valid_items):
            progress.progress(
                idx / len(valid_items),
                text=f"Article {idx + 1}/{len(valid_items)} : {item['keyword']}",
            )

            with log_container.status(
                f"Article {idx + 1} : {item['keyword']}", expanded=True
            ) as article_status:
                result = process_article(
                    url=item["url"],
                    keyword=item["keyword"],
                    api=api,
                    anthropic_key=anthropic_key,
                    sf_data=sf_data,
                    progress_callback=lambda msg: st.write(msg),
                )

                if result["success"]:
                    article_status.update(
                        label=f"✅ {item['keyword']}", state="complete", expanded=False
                    )
                else:
                    article_status.update(
                        label=f"❌ {item['keyword']}", state="error", expanded=False
                    )

                results.append(result)

        progress.progress(1.0, text="Terminé !")

        # Résumé
        st.divider()
        successful = [r for r in results if r["success"]]
        failed = [r for r in results if not r["success"]]

        st.success(f"**{len(successful)}/{len(results)}** articles générés")

        if failed:
            for r in failed:
                st.error(f"Échec : {r.get('url', '?')} — {r.get('error', '?')}")

        if successful:
            # ZIP
            zip_buffer = io.BytesIO()
            with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
                for r in successful:
                    slug = r["keyword"]
                    fname = f"{slug}.docx"
                    zf.writestr(fname, r["docx"].getvalue())

            zip_buffer.seek(0)

            st.download_button(
                f"📥 Télécharger les {len(successful)} articles (.zip)",
                data=zip_buffer,
                file_name="refresh_articles_gererseul.zip",
                mime="application/zip",
                use_container_width=True,
            )

            # Tableau récapitulatif
            st.subheader("📋 Résumé")
            df = pd.DataFrame([{
                "URL": r["url"],
                "Mot-clé": r["keyword"],
                "Title SEO": r["title_seo"],
                "Mots (avant)": r["existing_word_count"],
                "Mots (après)": r["new_word_count"],
            } for r in successful])
            st.dataframe(df, use_container_width=True, hide_index=True)

            st.caption(f"Coût DataForSEO total : ${api.total_cost:.4f}")

st.caption("🔄 Rafraîchissement de Contenu — Gérer Seul | Ma Toolbox SEO")
